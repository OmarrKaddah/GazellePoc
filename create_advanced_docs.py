"""
Generate additional mock banking documents designed to showcase graph retrieval advantages.

Key design principles:
1. INDIRECT references: Doc A mentions entity X, Doc B mentions entity X in a different context
   → vector search won't connect them, but graph will via shared entity nodes
2. CHAIN dependencies: A → B → C (you need to follow 2 hops to get the full answer)
3. SEMANTIC gaps: Two chunks discuss the same topic but use completely different vocabulary
4. HIGH chunk count: ~100+ chunks so top_k=5 only covers ~5% of corpus
5. CONDITIONAL logic split across documents

Run: python create_advanced_docs.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt


def create_aml_policy():
    """Anti-Money Laundering Policy - references: compliance committee, CRO, reporting thresholds."""
    doc = Document()
    doc.add_heading("Anti-Money Laundering and Counter-Terrorism Financing Policy", level=0)
    doc.add_paragraph("Document ID: POL-AML-2025-004")
    doc.add_paragraph("Effective Date: January 1, 2025")
    doc.add_paragraph("Classification: Strictly Confidential")

    doc.add_heading("1. Regulatory Framework", level=1)
    doc.add_paragraph(
        "This policy is issued in compliance with Central Bank Circular No. 2024/AML/012 and the "
        "Anti-Money Laundering Law No. 80 of 2002 (as amended). The Bank must maintain a risk-based "
        "approach to AML/CTF in accordance with FATF Recommendations."
    )

    doc.add_heading("2. Customer Due Diligence (CDD)", level=1)
    doc.add_heading("2.1 Standard CDD", level=2)
    doc.add_paragraph(
        "Standard CDD must be performed for all new customers. This includes identity verification, "
        "beneficial ownership identification, and understanding the purpose of the business relationship. "
        "For individual accounts, a valid national ID or passport is required along with proof of address "
        "dated within the last 3 months."
    )

    doc.add_heading("2.2 Enhanced Due Diligence (EDD)", level=2)
    doc.add_paragraph(
        "Enhanced Due Diligence is mandatory for: (a) Politically Exposed Persons (PEPs), "
        "(b) customers from high-risk jurisdictions as defined in Appendix A of the Country Risk Assessment, "
        "(c) correspondent banking relationships, and (d) any customer with transactions flagged by the "
        "Transaction Monitoring System (TMS). EDD requires approval from the Compliance Committee "
        "(see Governance Policy, Section 3.2 for committee composition)."
    )

    doc.add_heading("3. Transaction Monitoring", level=1)
    doc.add_heading("3.1 Monitoring Thresholds", level=2)
    doc.add_paragraph(
        "All cash transactions exceeding EGP 300,000 (or equivalent in foreign currency, calculated using "
        "the daily FX rates published by the Treasury Department as per FX Policy Section 3) must be reported "
        "to the Money Laundering Reporting Officer (MLRO). Multiple transactions by the same customer "
        "aggregating to EGP 500,000 within a 30-day rolling window must trigger an automatic alert."
    )

    doc.add_heading("3.2 Suspicious Activity Reporting", level=2)
    doc.add_paragraph(
        "IF a transaction is flagged as suspicious by the TMS, THEN the branch must file a preliminary "
        "report to the MLRO within 24 hours. IF the MLRO confirms the suspicion, THEN a Suspicious "
        "Transaction Report (STR) must be filed with the Financial Intelligence Unit (FIU) within 3 business "
        "days. The MLRO reports to the Chief Risk Officer (CRO) who has ultimate responsibility for "
        "AML compliance oversight (see Risk Governance Policy, Section 2.1)."
    )

    doc.add_heading("4. Sanctions Screening", level=1)
    doc.add_paragraph(
        "All customers and transactions must be screened against: UN Security Council sanctions lists, "
        "OFAC SDN list, EU sanctions lists, and local Central Bank restricted lists. Sanctions screening "
        "must occur at onboarding, periodically (at least quarterly), and for every wire transfer. "
        "Any positive match must freeze the transaction and be escalated to the Compliance Committee "
        "within 2 hours."
    )

    doc.add_heading("5. Record Retention", level=1)
    doc.add_paragraph(
        "All CDD records must be retained for a minimum of 5 years after the end of the customer "
        "relationship. Transaction records must be retained for 10 years. STR-related records must be "
        "retained indefinitely until cleared by the FIU. These requirements align with the Bank's "
        "Data Governance Policy (POL-DG-2025-009, Section 4)."
    )

    doc.add_heading("6. العقوبات والمخالفات", level=1)
    doc.add_paragraph(
        "أي موظف يخالف أحكام سياسة مكافحة غسل الأموال يتعرض لإجراءات تأديبية وفقاً لسياسة الموارد البشرية. "
        "في حالة المخالفات الجسيمة، يتم إبلاغ الجهات الرقابية المختصة. "
        "يتحمل مدير المخاطر التنفيذي (CRO) المسؤولية النهائية عن الامتثال لمتطلبات مكافحة غسل الأموال."
    )

    save_path = Path("data/raw_docs/AML_Policy_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_governance_policy():
    """Corporate Governance - defines committees that OTHER policies reference."""
    doc = Document()
    doc.add_heading("Corporate Governance and Committee Structure", level=0)
    doc.add_paragraph("Document ID: POL-GOV-2025-005")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Board of Directors", level=1)
    doc.add_paragraph(
        "The Board of Directors consists of 11 members including 4 independent directors. "
        "The Board meets at least quarterly and has ultimate oversight responsibility for all risk "
        "management activities. The Board delegates operational risk management to the following committees."
    )

    doc.add_heading("2. Risk Management Committees", level=1)
    doc.add_heading("2.1 Board Risk Committee (BRC)", level=2)
    doc.add_paragraph(
        "The Board Risk Committee consists of 5 Board members (minimum 3 independent). It meets monthly "
        "and reviews: aggregate risk exposure reports, risk appetite framework compliance, and capital "
        "adequacy assessments. The BRC has authority to impose temporary trading halts on any business unit "
        "when risk limits are breached (see FX Policy Section 2.2 and Credit Risk Policy Section 3)."
    )

    doc.add_heading("2.2 Executive Risk Committee (ERC)", level=2)
    doc.add_paragraph(
        "The Executive Risk Committee is chaired by the CEO and includes: CRO, CFO, Chief Compliance "
        "Officer (CCO), Head of Treasury, and Head of Credit. It meets weekly and oversees day-to-day risk "
        "management. The ERC approves risk limit changes up to 10% of existing limits; changes exceeding "
        "10% require BRC approval."
    )

    doc.add_heading("3. Specialized Committees", level=1)
    doc.add_heading("3.1 Credit Committee Structure", level=2)
    doc.add_paragraph(
        "Credit approval follows a tiered committee structure: Branch Credit Committee (up to EGP 5M), "
        "Regional Credit Committee (up to EGP 25M), Head Office Credit Committee (up to EGP 100M), "
        "Board of Directors (above EGP 100M). See Credit Risk Policy Section 2.1 for detailed approval "
        "matrix and conditions."
    )

    doc.add_heading("3.2 Compliance Committee", level=2)
    doc.add_paragraph(
        "The Compliance Committee is chaired by the CCO and includes: MLRO, Head of Legal, Head of "
        "Internal Audit, and a Board-appointed independent member. It meets bi-weekly and oversees: "
        "AML/CTF compliance (see AML Policy), sanctions screening, regulatory reporting, and customer "
        "complaints. The Compliance Committee reports to the BRC."
    )

    doc.add_heading("3.3 Treasury Committee", level=2)
    doc.add_paragraph(
        "The Treasury Committee is chaired by the Head of Treasury and includes: CRO, CFO, and Head of "
        "Market Risk. It meets daily for FX position reviews and weekly for ALM reviews. The Treasury "
        "Committee has authority to approve exotic currency pair trading (see FX Policy Section 3) "
        "and manages the Bank's liquidity buffer requirements."
    )

    doc.add_heading("4. Chief Risk Officer (CRO) Mandate", level=1)
    doc.add_paragraph(
        "The CRO has a dual reporting line: operationally to the CEO and functionally to the BRC Chair. "
        "The CRO's mandate includes: (a) oversight of credit, market, operational, and liquidity risk, "
        "(b) ultimate authority over risk limit setting subject to BRC approval, (c) AML/CTF compliance "
        "oversight (delegated from the Board), (d) quarterly risk appetite review, and "
        "(e) veto power on transactions exceeding 50% of any risk limit."
    )

    doc.add_heading("5. Escalation Chain", level=1)
    doc.add_paragraph(
        "Risk events must be escalated through the following chain: "
        "Line Manager → Department Head → relevant Committee → ERC → BRC → Board of Directors. "
        "Critical events (fraud, regulatory breach, loss >EGP 10M) skip directly to ERC and BRC. "
        "The CRO has authority to invoke emergency escalation at any level."
    )

    save_path = Path("data/raw_docs/Governance_Policy_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_liquidity_policy():
    """Liquidity Risk - references treasury, FX, capital adequacy."""
    doc = Document()
    doc.add_heading("Liquidity Risk Management Policy", level=0)
    doc.add_paragraph("Document ID: POL-LR-2025-006")
    doc.add_paragraph("Effective Date: February 1, 2025")

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This policy establishes the Bank's framework for managing liquidity risk to ensure the Bank "
        "can meet all financial obligations as they fall due under both normal and stressed conditions. "
        "It complements the Capital Adequacy Policy (POL-CA-2025-007) and must be read in conjunction "
        "with the FX Policy for foreign currency liquidity management."
    )

    doc.add_heading("2. Liquidity Ratios", level=1)
    doc.add_heading("2.1 Liquidity Coverage Ratio (LCR)", level=2)
    doc.add_paragraph(
        "The Bank must maintain a Liquidity Coverage Ratio (LCR) of at least 100% at all times. "
        "The LCR is calculated as: High Quality Liquid Assets (HQLA) / Total Net Cash Outflows over "
        "30 days. The Treasury Department (see Governance Policy, Section 3.3) is responsible for "
        "daily LCR monitoring and reporting."
    )

    doc.add_heading("2.2 Net Stable Funding Ratio (NSFR)", level=2)
    doc.add_paragraph(
        "The Net Stable Funding Ratio must be maintained above 100%. Available Stable Funding must "
        "exceed Required Stable Funding at all times. NSFR is reported monthly to the Board Risk "
        "Committee (see Governance Policy, Section 2.1)."
    )

    doc.add_heading("3. Foreign Currency Liquidity", level=1)
    doc.add_paragraph(
        "Foreign currency liquidity must be managed separately for each significant currency. "
        "The Bank must maintain a minimum FX liquidity buffer equal to 15% of FX liabilities. "
        "FX liquidity positions are subject to the Net Open Position limits defined in the FX Policy "
        "(Section 2.1). IF the FX liquidity buffer falls below 20%, THEN the Treasury Committee must "
        "convene an emergency meeting within 4 hours."
    )

    doc.add_heading("4. Stress Testing", level=1)
    doc.add_paragraph(
        "Liquidity stress tests must be conducted monthly using three scenarios: "
        "(a) idiosyncratic stress (bank-specific), (b) market-wide stress, (c) combined stress. "
        "Stress test results must demonstrate survival for at least 30 days under the combined scenario. "
        "IF any stress scenario shows a liquidity shortfall, THEN the Contingency Funding Plan must "
        "be activated (see Section 5) and the CRO must notify the Board Risk Committee."
    )

    doc.add_heading("5. Contingency Funding Plan (CFP)", level=1)
    doc.add_paragraph(
        "The CFP defines actions to be taken in a liquidity crisis: "
        "Stage 1 (Early Warning): LCR drops below 120% → Treasury Committee convenes, activates monitoring. "
        "Stage 2 (Elevated Risk): LCR drops below 110% → ERC notifies BRC, asset liquidation options reviewed. "
        "Stage 3 (Crisis): LCR drops below 100% → Board convenes, Central Bank notified, emergency "
        "asset sales authorized. Each stage requires specific credit facility adjustments as defined "
        "in the Credit Risk Policy (Section 3 - Concentration Limits)."
    )

    doc.add_heading("6. متطلبات السيولة بالعملة المحلية", level=1)
    doc.add_paragraph(
        "يجب على البنك الحفاظ على نسبة سيولة بالعملة المحلية لا تقل عن 20% من إجمالي الودائع بالجنيه المصري. "
        "يتم احتساب هذه النسبة يومياً وتقديمها إلى البنك المركزي أسبوعياً. "
        "في حالة انخفاض النسبة عن 25%، يتم تفعيل إجراءات الإنذار المبكر."
    )

    save_path = Path("data/raw_docs/Liquidity_Policy_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_capital_adequacy_policy():
    """Capital Adequacy - Tier 1 capital definition that ALL other policies reference."""
    doc = Document()
    doc.add_heading("Capital Adequacy Policy", level=0)
    doc.add_paragraph("Document ID: POL-CA-2025-007")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Capital Structure", level=1)
    doc.add_heading("1.1 Tier 1 Capital Definition", level=2)
    doc.add_paragraph(
        "Tier 1 Capital (Core Capital) consists of: Common Equity Tier 1 (CET1) capital comprising "
        "paid-up capital, retained earnings, and other comprehensive income reserves, net of regulatory "
        "deductions. As of December 31, 2024, the Bank's Tier 1 capital stands at EGP 2,500,000,000 "
        "(two billion five hundred million Egyptian Pounds). This figure is used as the denominator for "
        "all risk limit calculations across the Bank's policies."
    )

    doc.add_heading("1.2 Tier 2 Capital", level=2)
    doc.add_paragraph(
        "Tier 2 Capital includes: subordinated debt instruments with original maturity of at least 5 years, "
        "general provisions (up to 1.25% of risk-weighted assets), and revaluation reserves at a 45% "
        "discount. Total Tier 2 capital shall not exceed 100% of Tier 1 capital."
    )

    doc.add_heading("2. Minimum Capital Ratios", level=1)
    doc.add_paragraph(
        "The Bank must maintain the following minimum capital ratios at all times: "
        "Common Equity Tier 1 (CET1) ratio: 7.0% (including 2.5% conservation buffer). "
        "Tier 1 capital ratio: 8.5%. "
        "Total capital ratio: 12.5% (including 2.5% conservation buffer). "
        "These ratios are calculated using risk-weighted assets (RWA) under the standardized approach."
    )

    doc.add_heading("3. Capital Allocation for Risk Types", level=1)
    doc.add_heading("3.1 Credit Risk Capital", level=2)
    doc.add_paragraph(
        "Credit risk capital is allocated based on the standardized approach under Basel III. "
        "The single borrower limit of 20% of Tier 1 capital (see Credit Risk Policy, Section 3) "
        "translates to a maximum exposure of EGP 500,000,000 per borrower. Connected party limit of "
        "25% of Tier 1 capital equals EGP 625,000,000."
    )

    doc.add_heading("3.2 Market Risk Capital (FX)", level=2)
    doc.add_paragraph(
        "Market risk capital for FX positions is calculated using the standardized measurement method. "
        "The aggregate NOP limit of 25% of Tier 1 capital (see FX Policy, Section 2.1) equals "
        "EGP 625,000,000. Individual currency position limit of 15% equals EGP 375,000,000. "
        "Capital charge for FX risk is 8% of the higher of total long or total short positions."
    )

    doc.add_heading("3.3 Operational Risk Capital", level=2)
    doc.add_paragraph(
        "Operational risk capital is calculated using the Basic Indicator Approach: 15% of average "
        "gross income over the last 3 years. The Bank allocates an additional operational risk buffer "
        "of 2% of Tier 1 capital (EGP 50,000,000) for unexpected operational losses. "
        "This buffer is reviewed quarterly by the Executive Risk Committee."
    )

    doc.add_heading("4. Capital Conservation Actions", level=1)
    doc.add_paragraph(
        "IF the total capital ratio falls below 13.0%, THEN dividend distributions are restricted: "
        "the Bank may distribute a maximum of 60% of earnings. "
        "IF the ratio falls below 12.5%, THEN all dividends are suspended and a capital restoration plan "
        "must be submitted to the Central Bank within 30 days. "
        "IF the CET1 ratio falls below 5.5%, THEN the Bank enters Prompt Corrective Action (PCA) "
        "and must reduce risk-weighted assets by curtailing new lending (Credit Risk Policy) and "
        "closing proprietary trading positions (FX Policy, Section 2.2)."
    )

    doc.add_heading("5. العلاقة مع سياسات المخاطر الأخرى", level=1)
    doc.add_paragraph(
        "تستخدم جميع سياسات إدارة المخاطر رأس المال من الشريحة الأولى كأساس لحساب حدود المخاطر: "
        "سياسة مخاطر الائتمان: الحد الأقصى للمقترض الواحد = 20% = 500 مليون جنيه. "
        "سياسة الصرف الأجنبي: حد المركز المفتوح الصافي = 25% = 625 مليون جنيه. "
        "سياسة مخاطر التشغيل: احتياطي الخسائر غير المتوقعة = 2% = 50 مليون جنيه."
    )

    save_path = Path("data/raw_docs/Capital_Adequacy_Policy_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_it_security_policy():
    """IT Security - connects to operational risk, AML systems, data governance."""
    doc = Document()
    doc.add_heading("Information Technology and Cybersecurity Policy", level=0)
    doc.add_paragraph("Document ID: POL-IT-2025-008")
    doc.add_paragraph("Effective Date: March 1, 2025")

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This policy covers all information technology systems, cybersecurity controls, and digital "
        "assets of the Bank. It applies to all employees, contractors, and third-party service providers "
        "with access to Bank systems. The IT Department is the primary Control Owner for system-related "
        "operational risks (see Operational Risk Framework, Section 2)."
    )

    doc.add_heading("2. Core Banking System Security", level=1)
    doc.add_paragraph(
        "The Core Banking System (CBS) processes all customer transactions including deposits, loans, "
        "and FX transactions. Access to CBS is role-based with the following privilege levels: "
        "Level 1 (View Only): branch staff without approval authority. "
        "Level 2 (Transaction): staff authorized for transactions up to their approval limit as defined "
        "in the Credit Risk Policy (Section 2.1) and FX Policy (Section 2). "
        "Level 3 (Admin): IT administrators only, requires dual authorization."
    )

    doc.add_heading("3. Transaction Monitoring System", level=1)
    doc.add_paragraph(
        "The Transaction Monitoring System (TMS) is the Bank's primary AML detection tool. "
        "It processes all transactions in real-time and applies rule-based and ML-based detection algorithms. "
        "The TMS thresholds are configured as per AML Policy Section 3.1: cash transactions above "
        "EGP 300,000 and aggregate transactions above EGP 500,000 in 30 days. "
        "TMS alerts are routed to the MLRO through the Compliance workflow system."
    )

    doc.add_heading("4. Cybersecurity Incident Response", level=1)
    doc.add_heading("4.1 Classification", level=2)
    doc.add_paragraph(
        "Security incidents are classified into four severity levels: "
        "P1 (Critical): Active data breach or ransomware affecting production systems. "
        "P2 (High): Unauthorized access detected, potential data exposure. "
        "P3 (Medium): Malware detected on internal network, contained. "
        "P4 (Low): Policy violation, phishing attempt blocked."
    )

    doc.add_heading("4.2 Escalation and Notification", level=2)
    doc.add_paragraph(
        "P1 incidents must be reported to the CRO within 1 hour and to the Board Risk Committee within "
        "4 hours. P1 and P2 incidents affecting customer data must be reported to the Central Bank within "
        "24 hours as per Operational Risk Framework Section 3. IF a cybersecurity incident results in "
        "financial loss exceeding EGP 500,000, THEN it must also be reported through the operational "
        "risk incident process (Operational Risk Framework, Section 3)."
    )

    doc.add_heading("5. Third-Party Risk", level=1)
    doc.add_paragraph(
        "All third-party technology vendors must undergo security assessment before onboarding. "
        "Cloud service providers must maintain SOC 2 Type II certification. "
        "Third-party access to Bank data is governed by the Data Governance Policy (POL-DG-2025-009). "
        "Outsourcing of critical IT functions requires Board Risk Committee approval "
        "(see Governance Policy, Section 2.1)."
    )

    doc.add_heading("6. Business Continuity for IT Systems", level=1)
    doc.add_paragraph(
        "The Bank maintains a secondary data center with real-time replication. Recovery Time Objective "
        "(RTO) for critical systems: 4 hours. Recovery Point Objective (RPO): 1 hour. "
        "IT business continuity is a component of the overarching BCM plan owned by the BCM Team "
        "(see Operational Risk Framework, Section 2 - Business Continuity category). "
        "Annual DR tests are reported to the Board Risk Committee."
    )

    save_path = Path("data/raw_docs/IT_Security_Policy_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_data_governance_policy():
    """Data Governance - connects to AML retention, IT security, regulatory reporting."""
    doc = Document()
    doc.add_heading("Data Governance Policy", level=0)
    doc.add_paragraph("Document ID: POL-DG-2025-009")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Data Classification", level=1)
    doc.add_paragraph(
        "All Bank data must be classified into one of four categories: "
        "Public: annual reports, published rates. "
        "Internal: operational procedures, internal communications. "
        "Confidential: customer data, financial records, risk reports. "
        "Strictly Confidential: Board minutes, strategic plans, AML investigation files."
    )

    doc.add_heading("2. Data Ownership", level=1)
    doc.add_paragraph(
        "Each data domain has a designated Data Owner responsible for classification, access control, "
        "and quality. Key data owners: "
        "Customer Data: Head of Retail Banking. "
        "Financial Data: CFO. "
        "Risk Data: CRO. "
        "Compliance Data: Chief Compliance Officer (CCO). "
        "Data ownership assignments are reviewed annually by the Executive Risk Committee "
        "(see Governance Policy, Section 2.2)."
    )

    doc.add_heading("3. Data Quality Standards", level=1)
    doc.add_paragraph(
        "Risk-related data used in capital adequacy calculations (see Capital Adequacy Policy, Section 2) "
        "must meet the following quality standards: Completeness ≥ 99%, Accuracy ≥ 99.5%, "
        "Timeliness: updated within 1 business day. Data quality metrics are reported quarterly "
        "to the Board Risk Committee."
    )

    doc.add_heading("4. Data Retention and Disposal", level=1)
    doc.add_paragraph(
        "Data retention periods are aligned with regulatory requirements: "
        "Customer identity records: minimum 5 years after relationship end (as per AML Policy, Section 5). "
        "Transaction records: minimum 10 years (as per AML Policy, Section 5). "
        "Credit facility records: minimum 7 years after facility closure. "
        "Board and committee minutes: permanent retention. "
        "Disposal of confidential data requires shredding (physical) or secure erasure (digital) with "
        "certification from the IT Department."
    )

    doc.add_heading("5. Regulatory Reporting Data", level=1)
    doc.add_paragraph(
        "The Bank submits the following regulatory reports using data governed by this policy: "
        "Daily: Net Open Position report (FX Policy), Large Exposure report. "
        "Weekly: Liquidity ratios (Liquidity Policy, Section 2). "
        "Monthly: Capital Adequacy report (Capital Adequacy Policy), BRC risk summary. "
        "Quarterly: IFRS 9 provisioning report (Credit Risk Policy, Section 5). "
        "All regulatory reports must be signed off by both the Data Owner and the CRO."
    )

    doc.add_heading("6. حماية البيانات الشخصية", level=1)
    doc.add_paragraph(
        "يلتزم البنك بحماية البيانات الشخصية للعملاء وفقاً لقانون حماية البيانات الشخصية رقم 151 لسنة 2020. "
        "يجب الحصول على موافقة العميل قبل مشاركة بياناته مع أطراف ثالثة. "
        "يتحمل مسؤول حماية البيانات (DPO) مسؤولية ضمان الامتثال لمتطلبات الخصوصية."
    )

    save_path = Path("data/raw_docs/Data_Governance_Policy_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_internal_audit_charter():
    """Internal Audit - reviews all other policies, connects to committees."""
    doc = Document()
    doc.add_heading("Internal Audit Charter", level=0)
    doc.add_paragraph("Document ID: POL-IA-2025-010")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Independence and Authority", level=1)
    doc.add_paragraph(
        "The Internal Audit function reports functionally to the Board Audit Committee and "
        "administratively to the CEO. The Head of Internal Audit has unrestricted access to all Bank "
        "records, personnel, and physical properties. Internal Audit operates independently from "
        "management and the three lines of defense model."
    )

    doc.add_heading("2. Audit Scope", level=1)
    doc.add_heading("2.1 Credit Risk Audit", level=2)
    doc.add_paragraph(
        "Annual audit of credit risk management covering: loan origination compliance with the "
        "Credit Risk Policy (POL-CR-2025-002), approval authority adherence (Section 2.1), "
        "collateral valuation accuracy (Section 4.1), and IFRS 9 provisioning adequacy (Section 5). "
        "Findings are reported to the Board Risk Committee and the Head Office Credit Committee."
    )

    doc.add_heading("2.2 FX Operations Audit", level=2)
    doc.add_paragraph(
        "Quarterly audit of FX operations covering: compliance with exposure limits (FX Policy, "
        "Section 2), authorized currency pairs (FX Policy, Section 3), stop-loss monitoring "
        "(FX Policy, Section 2.2), and NOP reporting accuracy. "
        "FX audit results are shared with the Treasury Committee and the CRO."
    )

    doc.add_heading("2.3 AML/CTF Audit", level=2)
    doc.add_paragraph(
        "Semi-annual audit of AML controls covering: CDD/EDD compliance (AML Policy, Section 2), "
        "TMS configuration and effectiveness (AML Policy, Section 3), sanctions screening accuracy "
        "(AML Policy, Section 4), and STR filing timeliness. The MLRO must provide responses to "
        "all findings within 15 business days."
    )

    doc.add_heading("2.4 IT and Cybersecurity Audit", level=2)
    doc.add_paragraph(
        "Annual audit of IT controls covering: access management (IT Security Policy, Section 2), "
        "TMS system integrity (IT Security Policy, Section 3), incident response preparedness "
        "(IT Security Policy, Section 4), and DR testing results (IT Security Policy, Section 6). "
        "Critical IT findings are escalated directly to the Board Risk Committee."
    )

    doc.add_heading("3. Audit Committee Reporting", level=1)
    doc.add_paragraph(
        "The Head of Internal Audit presents a quarterly report to the Board Audit Committee covering: "
        "audit plan execution status, significant findings, management action plans, and overdue "
        "remediation items. High-risk findings that remain unresolved for more than 90 days are "
        "automatically escalated to the Board of Directors through the BRC (see Governance Policy, "
        "Section 5 - Escalation Chain)."
    )

    doc.add_heading("4. Coordination with External Audit", level=1)
    doc.add_paragraph(
        "Internal Audit coordinates with the external auditor to avoid duplication and ensure "
        "comprehensive coverage. The external auditor relies on Internal Audit work for: "
        "IFRS 9 provisioning testing, IT general controls testing, and regulatory compliance sampling. "
        "The external auditor's management letter findings are tracked alongside internal audit findings."
    )

    save_path = Path("data/raw_docs/Internal_Audit_Charter_2025.docx")
    doc.save(str(save_path))
    print(f"Created: {save_path}")


if __name__ == "__main__":
    print("Creating advanced mock banking documents...")
    create_aml_policy()
    create_governance_policy()
    create_liquidity_policy()
    create_capital_adequacy_policy()
    create_it_security_policy()
    create_data_governance_policy()
    create_internal_audit_charter()
    print(f"\nDone! 7 additional documents created in data/raw_docs/")
    print("\nCross-reference map:")
    print("  AML Policy → Governance (Compliance Committee), FX Policy (rates), Data Governance (retention)")
    print("  Governance → FX, Credit, AML, all policies (committee structure)")
    print("  Liquidity → FX (NOP), Capital Adequacy, Credit (concentration), Governance (BRC)")
    print("  Capital Adequacy → Credit, FX, Operational Risk (absolute limit calculations)")
    print("  IT Security → AML (TMS), Op Risk (incidents), Data Governance, Governance (BRC)")
    print("  Data Governance → AML (retention), Capital Adequacy (data quality), Governance (ERC)")
    print("  Internal Audit → Credit, FX, AML, IT, Governance (audits every other policy)")

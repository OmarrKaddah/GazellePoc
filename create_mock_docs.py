"""
Generate sample mock banking policy documents for testing.
Creates .docx files with realistic structure: headings, paragraphs, tables, Arabic + English.

Run: python create_mock_docs.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_fx_policy():
    """Document 1: Foreign Exchange Policy (English + Arabic)."""
    doc = Document()

    # Title
    title = doc.add_heading("Foreign Exchange Policy", level=0)
    doc.add_paragraph("Document ID: POL-FX-2025-001")
    doc.add_paragraph("Effective Date: January 1, 2025")
    doc.add_paragraph("Classification: Internal — Confidential")

    # Section 1
    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This policy establishes the framework for foreign exchange (FX) operations conducted by the Bank. "
        "It applies to all FX transactions including spot, forward, and swap transactions. "
        "All business units engaging in FX activities must comply with this policy."
    )
    doc.add_paragraph(
        "سياسة الصرف الأجنبي هذه تحدد الإطار العام لعمليات الصرف الأجنبي التي يجريها البنك. "
        "تنطبق على جميع معاملات الصرف الأجنبي بما في ذلك المعاملات الفورية والآجلة ومعاملات المقايضة."
    )

    # Section 2
    doc.add_heading("2. Exposure Limits", level=1)

    doc.add_heading("2.1 Net Open Position Limits", level=2)
    doc.add_paragraph(
        "The Bank's aggregate net open position (NOP) in all foreign currencies shall not exceed 25% of the Bank's Tier 1 capital. "
        "Individual currency positions shall not exceed 15% of Tier 1 capital. "
        "These limits are in accordance with Central Bank Circular No. 2024/FX/003."
    )

    # Table: Exposure limits
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ["Client Type", "Max Single Transaction (USD)", "Daily Limit (USD)", "NOP Limit (% of Tier 1)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["Corporate", "5,000,000", "20,000,000", "15%"],
        ["SME", "1,000,000", "5,000,000", "10%"],
        ["Retail", "100,000", "500,000", "5%"],
        ["Interbank", "50,000,000", "200,000,000", "25%"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph("")  # spacing

    doc.add_heading("2.2 Stop-Loss Limits", level=2)
    doc.add_paragraph(
        "A stop-loss limit of 2% of the position value shall be applied to all proprietary FX trading positions. "
        "If cumulative losses exceed 5% of allocated trading capital in any calendar month, "
        "all open proprietary positions must be closed and trading suspended pending management review. "
        "This requirement is governed by the Risk Management Framework (see Credit Risk Policy, Section 4.3)."
    )

    # Section 3
    doc.add_heading("3. Authorized Currencies", level=1)
    doc.add_paragraph(
        "The Bank is authorized to trade in the following currency pairs: USD/EGP, EUR/EGP, GBP/EGP, "
        "SAR/EGP, AED/EGP, JPY/EGP, CHF/EGP. Trading in exotic currency pairs requires prior approval "
        "from the Treasury Committee and must comply with Central Bank regulations."
    )

    # Arabic section
    doc.add_heading("4. المتطلبات التنظيمية", level=1)
    p = doc.add_paragraph(
        "يجب أن تتوافق جميع عمليات الصرف الأجنبي مع تعليمات البنك المركزي رقم 2024/FX/003. "
        "يتعين على البنك تقديم تقارير NOP يومية إلى البنك المركزي. "
        "في حالة تجاوز حد NOP بنسبة 25%، يجب إخطار البنك المركزي خلال 24 ساعة."
    )

    # Section 5
    doc.add_heading("5. Roles and Responsibilities", level=1)
    doc.add_paragraph(
        "The Chief Risk Officer (CRO) is responsible for monitoring FX exposure limits. "
        "The Treasury Department executes FX transactions within approved limits. "
        "Internal Audit conducts quarterly reviews of FX compliance. "
        "The Board Risk Committee receives monthly FX exposure reports."
    )

    save_path = Path("data/raw_docs/FX_Policy_2025.docx")
    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_credit_risk_policy():
    """Document 2: Credit Risk Policy (English + Arabic)."""
    doc = Document()

    title = doc.add_heading("Credit Risk Policy", level=0)
    doc.add_paragraph("Document ID: POL-CR-2025-002")
    doc.add_paragraph("Effective Date: January 1, 2025")

    # Section 1
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This policy defines the Bank's approach to credit risk management across all lending activities. "
        "It establishes approval authorities, concentration limits, and provisioning standards."
    )

    # Section 2
    doc.add_heading("2. Lending Authorities", level=1)

    doc.add_heading("2.1 Approval Matrix", level=2)
    doc.add_paragraph(
        "All credit facilities must be approved according to the following authority matrix. "
        "No single approver may approve a facility exceeding their delegated authority."
    )

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ["Authority Level", "Maximum Facility Amount (EGP)", "Conditions"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["Branch Manager", "5,000,000", "Secured facilities only"],
        ["Regional Credit Committee", "25,000,000", "Collateral coverage ≥ 120%"],
        ["Head Office Credit Committee", "100,000,000", "Full risk assessment required"],
        ["Board of Directors", "Above 100,000,000", "Special resolution required"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph("")

    # Section 3
    doc.add_heading("3. Concentration Limits", level=1)
    doc.add_paragraph(
        "Single borrower exposure shall not exceed 20% of the Bank's Tier 1 capital. "
        "Connected party exposure shall not exceed 25% of Tier 1 capital. "
        "Sector concentration shall not exceed 30% of total loan portfolio for any single sector."
    )

    doc.add_heading("3.1 Related Party Lending", level=2)
    doc.add_paragraph(
        "Loans to related parties are subject to additional restrictions as per Central Bank Circular No. 2024/CR/007. "
        "All related party facilities must be approved by the Board of Directors. "
        "The aggregate related party exposure shall not exceed 10% of Tier 1 capital."
    )

    # Section 4
    doc.add_heading("4. Collateral Requirements", level=1)

    doc.add_heading("4.1 Loan-to-Value Ratios", level=2)

    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    headers2 = ["Loan Type", "Maximum LTV Ratio", "Minimum Collateral Coverage"]
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h

    data2 = [
        ["Mortgage — Residential", "80%", "125%"],
        ["Mortgage — Commercial", "70%", "143%"],
        ["Working Capital (Secured)", "75%", "133%"],
        ["Project Finance", "65%", "154%"],
        ["Personal Loan (Secured)", "85%", "118%"],
    ]
    for r, row_data in enumerate(data2, 1):
        for c, val in enumerate(row_data):
            table2.rows[r].cells[c].text = val

    doc.add_paragraph("")

    doc.add_heading("4.2 Conditional Approval Criteria", level=2)
    doc.add_paragraph(
        "IF the borrower's debt-to-income ratio exceeds 40%, THEN the loan must be escalated to the Regional Credit Committee "
        "regardless of the facility amount. "
        "IF the borrower has existing exposure with the Bank exceeding EGP 10,000,000, THEN a comprehensive exposure review "
        "must be conducted before any new facility approval."
    )

    # Arabic section
    doc.add_heading("5. متطلبات التوفير", level=1)
    doc.add_paragraph(
        "يجب على البنك تطبيق معايير المعيار الدولي للتقارير المالية رقم 9 (IFRS 9) في تصنيف القروض وتكوين المخصصات. "
        "المرحلة الأولى: مخصص خسائر ائتمانية متوقعة لمدة 12 شهراً. "
        "المرحلة الثانية: مخصص خسائر ائتمانية متوقعة لمدى الحياة عند زيادة مخاطر الائتمان بشكل ملحوظ. "
        "المرحلة الثالثة: مخصص خسائر ائتمانية متوقعة لمدى الحياة للقروض المتعثرة."
    )

    # Cross-reference to FX policy
    doc.add_heading("6. FX-Denominated Credit Facilities", level=1)
    doc.add_paragraph(
        "Credit facilities denominated in foreign currencies are subject to both this Credit Risk Policy and the "
        "Foreign Exchange Policy (POL-FX-2025-001). The FX exposure arising from foreign currency loans must be "
        "included in the Bank's Net Open Position calculation (see FX Policy, Section 2.1). "
        "The combined credit and FX risk for any single borrower shall not exceed 30% of Tier 1 capital."
    )

    save_path = Path("data/raw_docs/Credit_Risk_Policy_2025.docx")
    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(save_path))
    print(f"Created: {save_path}")


def create_operational_risk_policy():
    """Document 3: Operational Risk Framework (English + Arabic)."""
    doc = Document()

    title = doc.add_heading("Operational Risk Management Framework", level=0)
    doc.add_paragraph("Document ID: POL-OR-2025-003")
    doc.add_paragraph("Effective Date: March 1, 2025")

    # Section 1
    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This framework covers all operational risks including process failures, system disruptions, "
        "human errors, fraud, and external events. It applies to all Bank departments and subsidiaries."
    )

    # Section 2
    doc.add_heading("2. Risk Categories", level=1)

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ["Category", "Risk Level", "Control Owner", "Review Frequency"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["Internal Fraud", "High", "Compliance Department", "Monthly"],
        ["External Fraud", "High", "Information Security", "Monthly"],
        ["System Failures", "Medium", "IT Department", "Quarterly"],
        ["Process Errors", "Medium", "Operations Department", "Quarterly"],
        ["Legal & Regulatory", "High", "Legal Department", "Monthly"],
        ["Business Continuity", "High", "BCM Team", "Semi-Annual"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph("")

    # Section 3
    doc.add_heading("3. Incident Escalation", level=1)
    doc.add_paragraph(
        "IF an operational loss exceeds EGP 500,000, THEN the incident must be reported to the Chief Risk Officer within 4 hours. "
        "IF an operational loss exceeds EGP 5,000,000, THEN the incident must be escalated to the Board Risk Committee within 24 hours. "
        "IF the incident involves potential fraud, THEN the Compliance Department and Internal Audit must be notified immediately "
        "regardless of the loss amount."
    )

    doc.add_heading("3.1 إجراءات التصعيد", level=2)
    doc.add_paragraph(
        "في حالة تجاوز الخسارة التشغيلية مبلغ 500,000 جنيه مصري، يجب إبلاغ مدير المخاطر التنفيذي خلال 4 ساعات. "
        "يجب تسجيل جميع الحوادث التشغيلية في نظام إدارة الحوادث خلال 24 ساعة من اكتشافها."
    )

    # Section 4 - cross-reference
    doc.add_heading("4. Integration with Credit and FX Risk", level=1)
    doc.add_paragraph(
        "Operational risk events affecting credit operations must also be assessed under the Credit Risk Policy "
        "(POL-CR-2025-002, Section 4.2). FX operational failures must be reported to the Treasury Department "
        "as per the FX Policy (POL-FX-2025-001, Section 5). "
        "The combined operational and credit loss for any single event shall trigger a comprehensive review "
        "if total losses exceed EGP 10,000,000."
    )

    save_path = Path("data/raw_docs/Operational_Risk_Framework_2025.docx")
    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(save_path))
    print(f"Created: {save_path}")


if __name__ == "__main__":
    print("Creating mock banking policy documents...")
    create_fx_policy()
    create_credit_risk_policy()
    create_operational_risk_policy()
    print("\nDone! Documents saved to data/raw_docs/")

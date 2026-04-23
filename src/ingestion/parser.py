"""
Document parser using Docling for structured extraction of Word documents.
Extracts sections, paragraphs, tables, and figures with full metadata.
"""

import json
import hashlib
import warnings
from pathlib import Path
from typing import Any, Optional
from dataclasses import dataclass, field, asdict

from docling.document_converter import DocumentConverter

# Module-level tracker for parse failures —
# lets build_index.py report skipped documents.
PARSE_FAILURES: list[tuple[str, str]] = []


@dataclass
class ParsedElement:
    """An atomic unit extracted from a document."""
    element_id: str
    doc_name: str
    doc_path: str
    element_type: str  # "paragraph", "table", "heading", "list_item", "figure"
    content: str
    section_path: list[str] = field(default_factory=list)  # ["Chapter 1", "Section 1.2", "1.2.1"]
    metadata: dict = field(default_factory=dict)
    language: str = "unknown"
    page_number: Optional[int] = None
    table_data: Optional[dict] = None  # Structured table if element_type == "table"

    def to_dict(self) -> dict:
        return asdict(self)


def _generate_element_id(doc_name: str, element_type: str, content: str, index: int) -> str:
    """Generate a deterministic unique ID for an element."""
    raw = f"{doc_name}::{element_type}::{index}::{content[:100]}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


def _detect_language(text: str) -> str:
    """Simple heuristic language detection (Arabic vs English)."""
    if not text.strip():
        return "unknown"
    arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F')
    ratio = arabic_chars / max(len(text), 1)
    if ratio > 0.3:
        return "ar"
    return "en"


def parse_document(doc_path: str | Path) -> list[ParsedElement]:
    """
    Parse a single Word document into structured elements.
    
    Returns a list of ParsedElement objects with full metadata.
    """
    doc_path = Path(doc_path)
    doc_name = doc_path.stem

    # Initialize Docling converter
    converter = DocumentConverter()

    # Convert document
    result = converter.convert(str(doc_path))
    doc = result.document

    elements: list[ParsedElement] = []
    current_section_path: list[str] = []
    element_index = 0

    # Pre-collect self_refs of every item that lives inside a table so the main
    # loop can skip them.  Docling's iterate_items() yields table-cell TextItems
    # as siblings of the TableItem, causing them to be double-emitted as
    # paragraphs unless we explicitly filter them out.
    table_owned_refs: set[str] = set()

    def _gather_table_refs(ref_or_item: Any) -> None:
        item_ = ref_or_item.resolve(doc=doc) if hasattr(ref_or_item, 'resolve') else ref_or_item
        ref = str(getattr(item_, 'self_ref', '') or '')
        if ref:
            table_owned_refs.add(ref)
        for child in getattr(item_, 'children', None) or []:
            _gather_table_refs(child)

    for tbl in getattr(doc, 'tables', []):
        for child in getattr(tbl, 'children', None) or []:
            _gather_table_refs(child)

    # Iterate over Docling's document items
    for item, _ in doc.iterate_items():
        item_type = type(item).__name__
        # Skip items that live inside a table cell — they are already captured
        # by export_to_markdown/export_to_dataframe on the TableItem itself.
        self_ref = str(getattr(item, 'self_ref', '') or '')
        if self_ref and self_ref in table_owned_refs:
            continue

        # BUG 1 FIX: determine table status once so the branches are mutually exclusive.
        is_table = item_type == "TableItem" or (
            hasattr(item, 'export_to_markdown') and 'table' in item_type.lower()
        )

        if not is_table and (
            item_type == "SectionHeaderItem"
            or (hasattr(item, 'label') and 'heading' in str(getattr(item, 'label', '')).lower())
        ):
            heading_text = item.text if hasattr(item, 'text') else str(item)
            heading_text = heading_text.strip()
            if heading_text:
                # BUG 2 FIX: coerce level to int; _level is hierarchy depth, not heading
                # level — never use it as a heading-level fallback.
                try:
                    level = int(getattr(item, 'level', None) or 1)
                except (TypeError, ValueError):
                    level = 1
                if level > 0:
                    current_section_path = current_section_path[:level - 1]
                current_section_path.append(heading_text)

                elem = ParsedElement(
                    element_id=_generate_element_id(doc_name, "heading", heading_text, element_index),
                    doc_name=doc_name,
                    doc_path=str(doc_path),
                    element_type="heading",
                    content=heading_text,
                    section_path=list(current_section_path),
                    language=_detect_language(heading_text),
                )
                elements.append(elem)
                element_index += 1

        elif is_table:
            # BUG 1 FIX: table branch is now elif — never double-counted as paragraph.
            try:
                # Pass doc so RichTableCell._get_text() can resolve references instead
                # of returning the placeholder "<!-- rich cell -->".
                table_md = item.export_to_markdown(doc=doc) if hasattr(item, 'export_to_markdown') else None
                if table_md:
                    table_data = None
                    if hasattr(item, 'export_to_dataframe'):
                        try:
                            df = item.export_to_dataframe(doc=doc)
                            table_data = {
                                "headers": [str(col) for col in df.columns],
                                "rows": [[str(cell) for cell in row] for row in df.values.tolist()],
                                "shape": list(df.shape),
                            }
                        except Exception as df_err:
                            # BUG 3 FIX: emit a warning instead of silently swallowing.
                            warnings.warn(
                                f"Table structured-data extraction failed in {doc_name} "
                                f"(section: {' > '.join(current_section_path) or 'root'}): {df_err}"
                            )
                    elem = ParsedElement(
                        element_id=_generate_element_id(doc_name, "table", table_md, element_index),
                        doc_name=doc_name,
                        doc_path=str(doc_path),
                        element_type="table",
                        content=table_md,
                        section_path=list(current_section_path),
                        language=_detect_language(table_md),
                        table_data=table_data,
                    )
                    elements.append(elem)
                    element_index += 1
            except Exception as e:
                warnings.warn(
                    f"Table extraction failed in {doc_name} "
                    f"(section: {' > '.join(current_section_path) or 'root'}): {e}"
                )

        elif hasattr(item, 'text') and item.text and item.text.strip():
            text = item.text.strip()
            elem = ParsedElement(
                element_id=_generate_element_id(doc_name, "paragraph", text, element_index),
                doc_name=doc_name,
                doc_path=str(doc_path),
                element_type="paragraph",
                content=text,
                section_path=list(current_section_path),
                language=_detect_language(text),
            )
            elements.append(elem)
            element_index += 1

    return elements


def parse_document_fallback(doc_path: str | Path) -> list[ParsedElement]:
    """
    Fallback parser using python-docx directly.
    Used when Docling has issues with specific document formats.
    """
    from docx import Document

    doc_path = Path(doc_path)
    doc_name = doc_path.stem
    doc = Document(str(doc_path))

    elements: list[ParsedElement] = []
    current_section_path: list[str] = []
    element_index = 0

    # BUG 4 FIX: build element-id → object maps so we can iterate body children
    # in document order and assign correct section context to each table.
    para_map: dict[int, Any] = {id(p._element): p for p in doc.paragraphs}
    table_map: dict[int, Any] = {id(t._element): t for t in doc.tables}

    for child in doc.element.body:
        child_id = id(child)

        if child_id in para_map:
            para = para_map[child_id]
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name.lower() if para.style else ""
            if "heading" in style_name:
                level = 1
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                current_section_path = current_section_path[:level - 1]
                current_section_path.append(text)
                elem = ParsedElement(
                    element_id=_generate_element_id(doc_name, "heading", text, element_index),
                    doc_name=doc_name,
                    doc_path=str(doc_path),
                    element_type="heading",
                    content=text,
                    section_path=list(current_section_path),
                    language=_detect_language(text),
                )
            else:
                elem = ParsedElement(
                    element_id=_generate_element_id(doc_name, "paragraph", text, element_index),
                    doc_name=doc_name,
                    doc_path=str(doc_path),
                    element_type="paragraph",
                    content=text,
                    section_path=list(current_section_path),
                    language=_detect_language(text),
                )
            elements.append(elem)
            element_index += 1

        elif child_id in table_map:
            table = table_map[child_id]
            headers: list[str] = []
            rows: list[list[str]] = []
            for i, row in enumerate(table.rows):
                # BUG 5 FIX: skip merged-cell duplicates via underlying XML element identity.
                seen_tcs: set[int] = set()
                cells: list[str] = []
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id not in seen_tcs:
                        seen_tcs.add(tc_id)
                        cells.append(cell.text.strip())
                if i == 0:
                    headers = cells
                else:
                    rows.append(cells)

            if not headers:
                continue

            md_lines = ["| " + " | ".join(headers) + " |"]
            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows:
                md_lines.append("| " + " | ".join(row) + " |")
            table_md = "\n".join(md_lines)
            elem = ParsedElement(
                element_id=_generate_element_id(doc_name, "table", table_md, element_index),
                doc_name=doc_name,
                doc_path=str(doc_path),
                element_type="table",
                content=table_md,
                section_path=list(current_section_path),
                language=_detect_language(table_md),
                table_data={"headers": headers, "rows": rows, "shape": [len(rows), len(headers)]},
            )
            elements.append(elem)
            element_index += 1

    return elements


def parse_all_documents(docs_dir: str | Path, output_dir: Optional[str | Path] = None) -> list[ParsedElement]:
    """
    Parse all Word documents in a directory.
    Optionally saves parsed output as JSON.
    """
    docs_dir = Path(docs_dir)
    PARSE_FAILURES.clear()  # Reset from previous runs
    all_elements: list[ParsedElement] = []

    doc_files = list(docs_dir.glob("*.docx")) + list(docs_dir.glob("*.doc"))
    print(f"Found {len(doc_files)} documents in {docs_dir}")
    print(doc_files)
    for doc_file in sorted(doc_files):
        if doc_file.name.startswith("~$"):
            continue  # Skip temp Word files
        print(f"  Parsing: {doc_file.name}")
        try:
            elements = parse_document(doc_file)
        except Exception as e:
            print(f"    Docling failed ({e}), trying fallback parser...")
            try:
                elements = parse_document_fallback(doc_file)
            except Exception as e2:
                warnings.warn(f"Both parsers failed for {doc_file.name}: {e2}. Skipping.")
                PARSE_FAILURES.append((doc_file.name, str(e2)))
                continue

        all_elements.extend(elements)
        print(f"    Extracted {len(elements)} elements")

    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / "parsed_elements.json"
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump([e.to_dict() for e in all_elements], f, ensure_ascii=False, indent=2)
        print(f"Saved {len(all_elements)} elements to {output_path}")

    # ── Report any parse failures ──
    if PARSE_FAILURES:
        warnings.warn(
            f"{len(PARSE_FAILURES)} document(s) failed to parse: "
            + ", ".join(name for name, _ in PARSE_FAILURES)
        )

    return all_elements

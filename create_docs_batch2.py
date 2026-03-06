"""
Generate 20 additional mock banking documents (docs 11-30) to create a larger corpus
where graph retrieval has a clear advantage over vector-only.

Design principles for graph advantage:
1. ENTITY CHAINS: Key answers require following entity links across 2-3 documents
2. VOCABULARY GAPS: Connected docs use different terminology for the same concepts
3. ROLE-RESPONSIBILITY SPLITS: "Who is responsible?" answers span multiple docs
4. THRESHOLD CASCADES: Numeric thresholds defined in one doc, consequences in another
5. REGULATORY CROSS-REFERENCES: Regulations cited in one doc, compliance procedures in another

Run: python create_docs_batch2.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

OUTPUT_DIR = Path("data/raw_docs")


def save(doc, name):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{name}.docx"
    doc.save(path)
    print(f"  Created {path}")


def doc_11_treasury_operations():
    doc = Document()
    doc.add_heading("Treasury Operations Manual", level=0)
    doc.add_paragraph("Document ID: MAN-TREAS-2025-011")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Treasury Department Structure", level=1)
    doc.add_paragraph(
        "The Treasury Department is headed by the Chief Treasury Officer (CTO), who reports directly "
        "to the Chief Financial Officer (CFO). The department comprises three desks: the Money Market Desk, "
        "the FX Trading Desk, and the Investment Securities Desk. Each desk is managed by a Senior Dealer "
        "who reports to the CTO."
    )

    doc.add_heading("2. Dealing Room Protocols", level=1)
    doc.add_heading("2.1 Pre-Trade Checks", level=2)
    doc.add_paragraph(
        "Before executing any trade, dealers must verify: (a) the counterparty limit via the Credit "
        "Administration System (CAS), (b) the instrument limit as defined in the Investment Policy "
        "(Document POL-INV-2025-018, Section 4.2), and (c) the aggregate position limit. The FX Trading "
        "Desk must additionally verify compliance with the Central Bank net open position (NOP) ceiling "
        "of 10% of Tier-1 capital as per the FX Policy."
    )

    doc.add_heading("2.2 Trade Execution", level=2)
    doc.add_paragraph(
        "All interbank FX trades exceeding USD 500,000 must be executed via the Bloomberg Terminal or "
        "Reuters Dealing. Trades below this threshold may be executed via telephone with mandatory voice "
        "recording. Every trade must be logged in the Treasury Management System (TMS-T) within 15 minutes "
        "of execution. The FX Trading Desk must report all daily positions to the Market Risk Unit "
        "by 16:00 local time."
    )

    doc.add_heading("3. Liquidity Buffer Management", level=1)
    doc.add_paragraph(
        "The Treasury Department maintains the High-Quality Liquid Assets (HQLA) portfolio to meet the "
        "Liquidity Coverage Ratio (LCR) requirement. The minimum LCR target is 120%, exceeding the "
        "regulatory minimum of 100% as specified in the Liquidity Policy (Section 2.2). The HQLA portfolio "
        "composition is reviewed weekly by the Asset and Liability Committee (ALCO)."
    )

    doc.add_heading("4. Funding Operations", level=1)
    doc.add_paragraph(
        "Daily funding requirements are determined by the Money Market Desk based on projected cash flows "
        "from the Operations Settlement Unit. Any funding shortfall exceeding EGP 500 million must be "
        "escalated to the CTO and the ALCO Chair. Emergency funding above EGP 1 billion triggers the "
        "Contingency Funding Plan (CFP) as defined in the Liquidity Policy, Section 5."
    )

    doc.add_heading("5. Counterparty Risk in Treasury", level=1)
    doc.add_paragraph(
        "Treasury counterparty exposure is monitored by the Credit Risk Unit using limits set by the "
        "Board Credit Committee. Pre-settlement Risk (PSR) for FX transactions is calculated at 3% of "
        "notional for spot and 6% for forwards. Any breach of counterparty limits must be reported to "
        "the Chief Risk Officer (CRO) within 1 hour and to the Board Risk Committee (BRC) at the next "
        "scheduled meeting."
    )

    save(doc, "Treasury_Operations_Manual_2025")


def doc_12_compliance_monitoring():
    doc = Document()
    doc.add_heading("Compliance Monitoring and Testing Program", level=0)
    doc.add_paragraph("Document ID: POL-COMP-2025-012")
    doc.add_paragraph("Effective Date: February 1, 2025")

    doc.add_heading("1. Program Objectives", level=1)
    doc.add_paragraph(
        "The Compliance Monitoring and Testing Program (CMTP) provides independent assurance that "
        "the Bank's activities comply with applicable laws, regulations, and internal policies. The "
        "program is owned by the Chief Compliance Officer (CCO) who reports to the Board Audit and "
        "Compliance Committee (BACC)."
    )

    doc.add_heading("2. Annual Compliance Plan", level=1)
    doc.add_heading("2.1 Risk-Based Prioritization", level=2)
    doc.add_paragraph(
        "The annual compliance testing plan is prioritized using the Compliance Risk Assessment (CRA) "
        "matrix. Areas rated High in the CRA must be tested at least semi-annually. Areas rated Medium "
        "must be tested annually. The CRA considers: regulatory change impact, prior findings severity, "
        "business volume, and operational complexity."
    )

    doc.add_heading("2.2 Testing Scope", level=2)
    doc.add_paragraph(
        "Mandatory testing areas include: (a) AML/CTF transaction monitoring effectiveness (per AML "
        "Policy Section 3), (b) sanctions screening accuracy, (c) consumer protection compliance, "
        "(d) data privacy adherence (per Data Governance Policy), and (e) market conduct in FX and "
        "fixed income trading (per FX Policy and Treasury Operations Manual)."
    )

    doc.add_heading("3. Testing Methodology", level=1)
    doc.add_paragraph(
        "Compliance testing employs three methods: (1) Transaction testing — random sampling of "
        "transactions to verify policy adherence, minimum sample size of 30 or 10% of population "
        "(whichever is larger); (2) Process walkthroughs — end-to-end review of key processes; "
        "(3) Control testing — evaluation of control design and operating effectiveness."
    )

    doc.add_heading("4. Findings and Remediation", level=1)
    doc.add_heading("4.1 Finding Classification", level=2)
    doc.add_paragraph(
        "Findings are classified as: Critical (regulatory breach with potential penalty), Major "
        "(significant policy violation), Minor (procedural gap with low risk). Critical findings "
        "must be reported to the BACC Chair within 24 hours and to the Central Bank Supervision "
        "Department within 5 business days if required by regulation."
    )

    doc.add_heading("4.2 Remediation Tracking", level=2)
    doc.add_paragraph(
        "All findings are logged in the Governance, Risk, and Compliance (GRC) system. Remediation "
        "deadlines are: Critical — 30 days, Major — 60 days, Minor — 90 days. The CCO provides a "
        "monthly remediation status report to the BACC. Any overdue Critical finding must be escalated "
        "to the Board of Directors."
    )

    doc.add_heading("5. Regulatory Examination Support", level=1)
    doc.add_paragraph(
        "The Compliance Department coordinates all regulatory examinations. Pre-examination preparation "
        "includes: assembling requested documents, briefing senior management, and identifying potential "
        "areas of concern based on the latest CRA. Post-examination, the CCO ensures all regulatory "
        "recommendations are entered into the GRC system with assigned owners from relevant business units."
    )

    save(doc, "Compliance_Monitoring_Program_2025")


def doc_13_hr_policy():
    doc = Document()
    doc.add_heading("Human Resources and Training Policy", level=0)
    doc.add_paragraph("Document ID: POL-HR-2025-013")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Mandatory Training Requirements", level=1)
    doc.add_paragraph(
        "All employees must complete the following annual training modules: (a) AML/CTF awareness "
        "(aligned with AML Policy Section 2 and 3), (b) Information Security awareness (per IT Security "
        "Policy Section 3), (c) Code of Conduct and Ethics, (d) Data Privacy (per Data Governance Policy "
        "Section 4). Training completion is tracked in the HR Management System (HRMS) and reported "
        "quarterly to the Board Governance Committee."
    )

    doc.add_heading("2. Specialized Training", level=1)
    doc.add_heading("2.1 Treasury and Trading Staff", level=2)
    doc.add_paragraph(
        "Staff assigned to the Treasury Department (including the FX Trading Desk, Money Market Desk, "
        "and Investment Securities Desk) must complete: (a) Market risk management certification within "
        "6 months of appointment, (b) Annual refresher on FX regulations and NOP limits, (c) Bloomberg "
        "Terminal proficiency assessment. The Chief Treasury Officer certifies completion to HR."
    )

    doc.add_heading("2.2 Risk and Compliance Staff", level=2)
    doc.add_paragraph(
        "Risk and compliance professionals must maintain relevant certifications (FRM, CFA, CAMS, or "
        "equivalent). The Bank sponsors certification fees. Staff failing to maintain certifications "
        "within 12 months of expiry are reassigned per the Performance Management framework."
    )

    doc.add_heading("3. Fit and Proper Requirements", level=1)
    doc.add_paragraph(
        "Senior management positions (C-suite, department heads, committee members) are subject to "
        "Central Bank fit-and-proper requirements. Background checks include: criminal record verification, "
        "credit history review, reference checks from previous employers, and verification of "
        "educational credentials. The Board Nomination Committee approves all senior appointments."
    )

    doc.add_heading("4. Succession Planning", level=1)
    doc.add_paragraph(
        "The Board Governance Committee oversees succession planning for all Key Management Personnel "
        "(KMP). KMP includes: CEO, CFO, CRO, CCO, CTO, Head of Internal Audit, and MLRO. Each KMP "
        "position must have at least two identified successors who receive targeted development. "
        "Succession plans are reviewed annually and reported to the Board of Directors."
    )

    doc.add_heading("5. Conflict of Interest Management", level=1)
    doc.add_paragraph(
        "Employees must disclose potential conflicts of interest annually via the COI Declaration Form "
        "in the HRMS. Trading desk employees are prohibited from personal trading in instruments "
        "covered by their desk without prior written approval from the CCO and their line manager. "
        "Violations are subject to disciplinary action per Section 8 of the Employee Handbook."
    )

    save(doc, "HR_Training_Policy_2025")


def doc_14_vendor_management():
    doc = Document()
    doc.add_heading("Third-Party and Vendor Risk Management Policy", level=0)
    doc.add_paragraph("Document ID: POL-VRM-2025-014")
    doc.add_paragraph("Effective Date: March 1, 2025")

    doc.add_heading("1. Scope and Applicability", level=1)
    doc.add_paragraph(
        "This policy governs the assessment, onboarding, and ongoing monitoring of all third-party "
        "vendors providing services to the Bank. Critical vendors (those supporting core banking, "
        "payment systems, cybersecurity, or cloud infrastructure) are subject to enhanced requirements. "
        "The Procurement Department manages the vendor lifecycle in coordination with the IT Security "
        "Department and the Operational Risk Unit."
    )

    doc.add_heading("2. Vendor Risk Assessment", level=1)
    doc.add_heading("2.1 Due Diligence Requirements", level=2)
    doc.add_paragraph(
        "Pre-onboarding due diligence includes: financial stability review, information security "
        "assessment (alignment with the Bank's IT Security Policy, Section 2), business continuity "
        "capability verification, regulatory compliance check, and reference verification. For vendors "
        "processing customer data, a Data Processing Impact Assessment (DPIA) must be completed per "
        "the Data Governance Policy, Section 6."
    )

    doc.add_heading("2.2 Risk Tiering", level=2)
    doc.add_paragraph(
        "Vendors are classified into three tiers: Tier-1 (Critical) — vendors whose failure would "
        "cause significant operational disruption, Tier-2 (Important) — vendors supporting significant "
        "business functions, Tier-3 (Standard) — all other vendors. The Business Continuity Management "
        "(BCM) Policy Section 4 defines the recovery requirements for each vendor tier."
    )

    doc.add_heading("3. Contract Requirements", level=1)
    doc.add_paragraph(
        "All vendor contracts must include: service level agreements (SLAs), data protection clauses "
        "(per Data Governance Policy), right-to-audit clauses, incident notification requirements "
        "(within 24 hours for security incidents per IT Security Policy Section 4.2), business continuity "
        "provisions, and exit/transition clauses. Legal Department approval is required for all Tier-1 "
        "vendor contracts."
    )

    doc.add_heading("4. Ongoing Monitoring", level=1)
    doc.add_paragraph(
        "Tier-1 vendors are reviewed semi-annually, Tier-2 annually, and Tier-3 biennially. Reviews "
        "assess: SLA performance, incident history, financial health, and control environment changes. "
        "The Vendor Risk Dashboard is reported to the Operational Risk Committee (ORC) quarterly. "
        "Concentration risk (over-reliance on single vendors) is monitored by the ORC with thresholds "
        "defined in the Operational Risk Framework, Section 6."
    )

    doc.add_heading("5. Cloud Service Providers", level=1)
    doc.add_paragraph(
        "Cloud service providers are automatically classified as Tier-1 and require additional approvals: "
        "(a) CRO sign-off on risk assessment, (b) CTO sign-off on technical architecture, (c) CCO "
        "confirmation of regulatory compliance (Central Bank cloud computing circular 2023/IT/003), "
        "(d) Board IT Committee approval for any workloads involving customer data or core banking functions."
    )

    save(doc, "Vendor_Risk_Management_2025")


def doc_15_bcm_policy():
    doc = Document()
    doc.add_heading("Business Continuity Management Policy", level=0)
    doc.add_paragraph("Document ID: POL-BCM-2025-015")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. BIA and Recovery Objectives", level=1)
    doc.add_paragraph(
        "The Business Impact Analysis (BIA) identifies critical business functions and their Recovery "
        "Time Objectives (RTO) and Recovery Point Objectives (RPO). Critical functions include: "
        "payment processing (RTO: 4 hours, RPO: 0), core banking system (RTO: 4 hours, RPO: 1 hour), "
        "FX trading (RTO: 2 hours, RPO: 0), ATM/POS network (RTO: 8 hours), and SWIFT messaging "
        "(RTO: 2 hours, RPO: 0)."
    )

    doc.add_heading("2. Crisis Management Team", level=1)
    doc.add_paragraph(
        "The Crisis Management Team (CMT) is chaired by the CEO and includes: CFO, CRO, CTO, CCO, "
        "Head of Operations, Head of IT, and the Corporate Communications Director. The CMT activates "
        "when a disruption is expected to exceed 4 hours or affects multiple business units. The CMT "
        "reports situation updates to the Board Chairman every 6 hours during active incidents."
    )

    doc.add_heading("3. IT Disaster Recovery", level=1)
    doc.add_paragraph(
        "The IT Disaster Recovery (DR) plan is maintained by the IT Department in coordination with "
        "the Business Continuity Officer. The primary data center is located in Cairo, with the DR "
        "site in Alexandria. Data replication is synchronous for core banking and FX trading systems, "
        "and asynchronous (max 1-hour lag) for all other systems. DR testing is conducted semi-annually "
        "with results reported to the Board IT Committee."
    )

    doc.add_heading("4. Pandemic and Staff Unavailability", level=1)
    doc.add_paragraph(
        "In case of staff unavailability exceeding 30%, the Bank activates split-team operations. "
        "Critical functions (as defined in Section 1) maintain minimum staffing levels: Treasury "
        "operations require at least 3 dealers (per Treasury Operations Manual, dealing room minimum "
        "staffing), payment processing requires 5 operators, and IT support requires 4 engineers. "
        "Remote work capability must be pre-tested quarterly."
    )

    doc.add_heading("5. Communication During Disruption", level=1)
    doc.add_paragraph(
        "External communications during business disruption events are managed exclusively by the "
        "Corporate Communications Director in consultation with the CEO and CCO. No employee may "
        "communicate with media or regulators about the disruption without CMT authorization. "
        "Regulatory notification (to Central Bank Supervision Department) must occur within 2 hours "
        "of CMT activation for events affecting payment systems or customer data, as per "
        "the Operational Risk Framework Section 3."
    )

    save(doc, "BCM_Policy_2025")


def doc_16_investment_policy():
    doc = Document()
    doc.add_heading("Investment Policy and Portfolio Management", level=0)
    doc.add_paragraph("Document ID: POL-INV-2025-016")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Investment Governance", level=1)
    doc.add_paragraph(
        "The Bank's investment portfolio is governed by the Asset and Liability Committee (ALCO), "
        "which sets the strategic asset allocation. The ALCO meets monthly and is chaired by the CEO "
        "with membership including CFO, CRO, CTO (Chief Treasury Officer), and Head of Risk Analytics. "
        "Investment decisions exceeding USD 50 million require Board approval via the Board "
        "Finance Committee."
    )

    doc.add_heading("2. Eligible Instruments", level=1)
    doc.add_paragraph(
        "The Bank may invest in: (a) Egyptian government securities (T-bills, T-bonds), (b) Central "
        "Bank certificates of deposit, (c) Investment-grade corporate bonds (minimum rating BBB- from "
        "Fitch/S&P or Baa3 from Moody's), (d) Listed equities (maximum 10% of portfolio), "
        "(e) Sukuk and Islamic financial instruments for the Islamic banking window. "
        "Prohibited instruments include: derivatives for speculation, cryptocurrency, unrated securities, "
        "and structured products without Board Finance Committee approval."
    )

    doc.add_heading("3. Concentration Limits", level=1)
    doc.add_paragraph(
        "Single issuer exposure must not exceed 15% of Tier-1 capital (excluding sovereign). Sector "
        "concentration must not exceed 25% of the investment portfolio. Country risk limits for "
        "non-domestic investments are set by the Board Risk Committee based on the Country Risk "
        "Assessment maintained by the Credit Risk Unit (per Credit Risk Policy, Section 7)."
    )

    doc.add_heading("4. Portfolio Limits and Thresholds", level=1)
    doc.add_heading("4.1 Duration Limits", level=2)
    doc.add_paragraph(
        "The average portfolio duration must remain within +/- 1 year of the benchmark set by ALCO. "
        "Individual security duration must not exceed 10 years. Duration deviations exceeding +/- 0.5 "
        "years from the benchmark require CTO approval; deviations exceeding +/- 1 year require ALCO approval."
    )

    doc.add_heading("4.2 Loss Limits", level=2)
    doc.add_paragraph(
        "Mark-to-market losses on the trading book exceeding 2% of portfolio value in any month "
        "trigger a mandatory review by the Market Risk Unit. Losses exceeding 5% require ALCO "
        "emergency meeting. Losses exceeding 8% trigger the Investment Stop-Loss Protocol: all new "
        "purchases are frozen and existing positions may only be unwound, subject to CFO and CRO approval. "
        "The Board Risk Committee is notified within 24 hours of any stop-loss activation."
    )

    doc.add_heading("5. Performance Measurement", level=1)
    doc.add_paragraph(
        "Investment performance is benchmarked against: Egyptian T-bill yield curve for local currency "
        "fixed income, and LIBOR/SOFR + spread for FCY investments. Monthly attribution reports are "
        "prepared by the Risk Analytics Unit and presented to ALCO. Annual portfolio performance is "
        "reported to the Board Finance Committee."
    )

    save(doc, "Investment_Policy_2025")


def doc_17_consumer_protection():
    doc = Document()
    doc.add_heading("Consumer Protection and Fair Dealing Policy", level=0)
    doc.add_paragraph("Document ID: POL-CP-2025-017")
    doc.add_paragraph("Effective Date: March 1, 2025")

    doc.add_heading("1. Fair Dealing Principles", level=1)
    doc.add_paragraph(
        "The Bank adheres to the Central Bank's Consumer Protection Regulations (Circular 2024/CP/001). "
        "All product pricing must be transparent with annual percentage rates (APR) clearly disclosed. "
        "No hidden fees or charges are permitted. The Compliance Department conducts quarterly mystery "
        "shopping exercises to verify branch-level adherence."
    )

    doc.add_heading("2. Complaint Handling", level=1)
    doc.add_heading("2.1 Internal Resolution", level=2)
    doc.add_paragraph(
        "Customer complaints must be acknowledged within 2 business days and resolved within 15 business "
        "days. Complex complaints involving multiple departments may extend to 30 business days with "
        "customer notification. The Customer Experience Unit tracks all complaints in the CRM system. "
        "Complaints involving potential regulatory violations must be copied to the CCO."
    )

    doc.add_heading("2.2 Escalation", level=2)
    doc.add_paragraph(
        "If a complaint is not resolved within 30 business days, or if the customer is dissatisfied "
        "with the resolution, it is escalated to the Senior Customer Ombudsman. Complaints involving "
        "fraud, unauthorized transactions, or discrimination are immediately escalated to the CCO and "
        "the Operational Risk Unit. Systemic complaint patterns (more than 10 similar complaints in "
        "a quarter) must be reported to the Board Governance Committee."
    )

    doc.add_heading("3. Product Suitability", level=1)
    doc.add_paragraph(
        "Before offering investment products to retail customers, the Bank must conduct a suitability "
        "assessment. This includes risk appetite questionnaire, financial literacy evaluation, and "
        "income/asset verification. Products classified as 'Complex' (including FX margin trading, "
        "structured deposits, and subordinated debt) may only be offered to customers classified as "
        "'Qualified' per the Central Bank classification criteria."
    )

    doc.add_heading("4. Vulnerable Customers", level=1)
    doc.add_paragraph(
        "The Bank maintains enhanced protections for vulnerable customers including: elderly (above 65), "
        "persons with disabilities, financially illiterate individuals (as identified by the literacy "
        "evaluation), and recent bereavement cases. Branch Managers must ensure vulnerable customers "
        "receive additional time and explanation. FX margin trading products are prohibited for "
        "vulnerable customers regardless of qualification status."
    )

    doc.add_heading("5. Fee Transparency and Disclosure", level=1)
    doc.add_paragraph(
        "All fees and charges must be published on the Bank's website and displayed in branches. "
        "FX conversion spreads must be disclosed before execution (per FX Policy, Section 6). "
        "Account maintenance fees, ATM fees, and wire transfer fees must be provided to the customer "
        "in a standardized Key Facts Statement (KFS) at account opening. The KFS template is approved "
        "by the CCO and reviewed annually by the BACC."
    )

    save(doc, "Consumer_Protection_Policy_2025")


def doc_18_market_risk():
    doc = Document()
    doc.add_heading("Market Risk Management Framework", level=0)
    doc.add_paragraph("Document ID: FRM-MKT-2025-018")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Market Risk Organization", level=1)
    doc.add_paragraph(
        "The Market Risk Unit (MRU) operates within the Risk Management Department, reporting to the "
        "CRO. The MRU is independent from the Treasury Department and has authority to set and monitor "
        "all market risk limits. The Head of MRU has direct access to the Board Risk Committee (BRC) "
        "for escalation purposes."
    )

    doc.add_heading("2. Value at Risk (VaR) Framework", level=1)
    doc.add_heading("2.1 VaR Methodology", level=2)
    doc.add_paragraph(
        "The Bank uses Historical Simulation VaR with a 99% confidence level and 1-day holding period "
        "for daily risk monitoring. A 10-day VaR is calculated for regulatory reporting. The observation "
        "window is 250 business days. Stressed VaR uses the 2008-2009 financial crisis period and the "
        "2023 Egyptian pound devaluation period."
    )

    doc.add_heading("2.2 VaR Limits", level=2)
    doc.add_paragraph(
        "VaR limits are set by the BRC annually. The aggregate 1-day VaR limit is USD 5 million. "
        "Desk-level limits are: FX Trading Desk — USD 2 million, Fixed Income Desk — USD 2.5 million, "
        "Equities Desk — USD 0.5 million. Any VaR utilization exceeding 80% of the limit triggers "
        "an amber alert to the CRO. Utilization exceeding 100% requires immediate position reduction "
        "and CRO approval to maintain positions."
    )

    doc.add_heading("3. FX Risk Specifics", level=1)
    doc.add_paragraph(
        "FX risk is managed through: (a) Net Open Position limits per currency pair (as defined in "
        "FX Policy, Section 4), (b) Intraday and overnight position limits, (c) Stop-loss triggers "
        "at 2% of position value. The daily FX risk report includes: NOP by currency, VaR contribution "
        "by FX desk, stress test results, and P&L attribution. This report is sent to: CRO, CTO, "
        "Treasury Head, and ALCO members by 17:00 daily."
    )

    doc.add_heading("4. Interest Rate Risk in the Banking Book (IRRBB)", level=1)
    doc.add_paragraph(
        "IRRBB is measured using: (a) Net Interest Income (NII) sensitivity to +/- 200bps shock, "
        "(b) Economic Value of Equity (EVE) sensitivity to +/- 200bps shock. The Board-approved limits "
        "are: NII impact must not exceed 10% of projected annual NII, EVE impact must not exceed 15% "
        "of Tier-1 capital. IRRBB results are reported to ALCO monthly and BRC quarterly."
    )

    doc.add_heading("5. Stress Testing", level=1)
    doc.add_paragraph(
        "Market risk stress testing is conducted monthly for desk-level scenarios and quarterly for "
        "bank-wide scenarios. Bank-wide scenarios include: (a) Egyptian pound devaluation of 20%, "
        "(b) Interest rate rise of 500bps, (c) Equity market decline of 30%, (d) Combined stress. "
        "Stress test results are reviewed by the CRO and presented to the BRC. If any stress scenario "
        "shows losses exceeding 20% of the capital buffer, the BRC must approve a remediation plan "
        "within 30 days."
    )

    save(doc, "Market_Risk_Framework_2025")


def doc_19_payment_systems():
    doc = Document()
    doc.add_heading("Payment Systems and Settlement Policy", level=0)
    doc.add_paragraph("Document ID: POL-PAY-2025-019")
    doc.add_paragraph("Effective Date: February 1, 2025")

    doc.add_heading("1. Payment Systems Overview", level=1)
    doc.add_paragraph(
        "The Bank participates in the following payment systems: RTGS (Real-Time Gross Settlement) "
        "operated by the Central Bank, ACH (Automated Clearing House) for batch payments, SWIFT for "
        "international wire transfers, and the National Payment Switch for card transactions. "
        "The Payment Operations Department manages daily settlement under the oversight of the "
        "Head of Operations who reports to the CFO."
    )

    doc.add_heading("2. Settlement Risk Management", level=1)
    doc.add_paragraph(
        "Pre-funding requirements for RTGS: the Bank must maintain a minimum balance of EGP 200 million "
        "in its Central Bank settlement account. Intraday liquidity monitoring is performed by the "
        "Treasury Money Market Desk (per Treasury Operations Manual, Section 4). Failed settlements "
        "exceeding EGP 50 million must be reported to the CRO and the Operational Risk Committee within "
        "2 hours."
    )

    doc.add_heading("3. SWIFT Operations", level=1)
    doc.add_paragraph(
        "SWIFT message types MT103 (customer transfers), MT202 (bank transfers), and MT300 (FX "
        "confirmations) are processed by the SWIFT Operations Unit. All SWIFT messages undergo "
        "sanctions screening before release (per AML Policy, Section 4). The Bank has adopted the "
        "SWIFT Customer Security Programme (CSP) and achieves annual attestation. CSP compliance "
        "is audited by Internal Audit (per Internal Audit Charter, Section 3)."
    )

    doc.add_heading("4. Card Payment Operations", level=1)
    doc.add_paragraph(
        "The Bank is PCI-DSS Level 1 certified. Card transaction processing is outsourced to "
        "PayTech Solutions Ltd. (a Tier-1 vendor per Vendor Risk Management Policy, Section 2.2). "
        "Daily card settlement reconciliation must be completed by 10:00 the following business day. "
        "Chargebacks exceeding USD 10,000 require Customer Experience Unit involvement per the "
        "Consumer Protection Policy, Section 2."
    )

    doc.add_heading("5. Correspondent Banking", level=1)
    doc.add_paragraph(
        "Correspondent banking relationships are subject to enhanced due diligence per AML Policy "
        "Section 2.2. Annual reviews include: KYC refresh, transaction volume analysis, and nested "
        "correspondent risk assessment. The Correspondent Banking Unit reports to the Head of "
        "Trade Finance who reports to the CFO. Any new correspondent relationship requires "
        "Board Credit Committee approval."
    )

    save(doc, "Payment_Systems_Policy_2025")


def doc_20_model_risk():
    doc = Document()
    doc.add_heading("Model Risk Management Policy", level=0)
    doc.add_paragraph("Document ID: POL-MRM-2025-020")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Model Governance", level=1)
    doc.add_paragraph(
        "The Bank's model inventory includes: credit scoring models, market risk models (VaR, as "
        "described in Market Risk Framework Section 2), operational risk models, ALM models, and "
        "AML transaction monitoring models. The Model Risk Committee (MRC), chaired by the CRO, "
        "oversees model development, validation, and retirement. The MRC reports to the Board "
        "Risk Committee."
    )

    doc.add_heading("2. Model Development Standards", level=1)
    doc.add_paragraph(
        "All models must be documented following the Model Documentation Template (MDT). Required "
        "elements include: purpose and scope, data requirements and quality assessment, methodology "
        "description, assumptions and limitations, validation results, and ongoing monitoring plan. "
        "Models must be developed independently from the business unit that uses them."
    )

    doc.add_heading("3. Model Validation", level=1)
    doc.add_paragraph(
        "Independent model validation is performed by the Model Validation Unit within the Risk "
        "Management Department. Validation must occur: (a) before initial deployment, (b) annually "
        "for all Tier-1 models (credit scoring, VaR, IFRS 9 ECL), (c) upon material changes to "
        "model inputs or methodology. Validation reports are submitted to the MRC and archived "
        "in the GRC system per the Internal Audit Charter requirements."
    )

    doc.add_heading("4. AI and Machine Learning Models", level=1)
    doc.add_heading("4.1 Additional Requirements", level=2)
    doc.add_paragraph(
        "AI/ML models are subject to enhanced governance: (a) Explainability assessment — models used "
        "for credit decisions must produce human-interpretable explanations, (b) Bias testing — "
        "models must be tested for demographic bias before deployment, (c) Drift monitoring — model "
        "performance is monitored monthly against baseline metrics, (d) Retraining governance — "
        "any retraining requires MRC approval."
    )

    doc.add_heading("4.2 Chatbot and NLP Models", level=2)
    doc.add_paragraph(
        "Customer-facing AI models (chatbots, virtual assistants) must: (a) clearly identify "
        "themselves as AI to customers, (b) escalate to human agents when confidence drops below "
        "the threshold defined in the Customer Experience Unit guidelines, (c) not provide investment "
        "advice or product recommendations without suitability validation (per Consumer Protection "
        "Policy, Section 3), (d) log all interactions for audit trail purposes per Data Governance "
        "Policy retention requirements."
    )

    doc.add_heading("5. Model Risk Reporting", level=1)
    doc.add_paragraph(
        "The CRO presents a quarterly Model Risk Dashboard to the BRC. The dashboard includes: "
        "model inventory summary, validation status, performance monitoring alerts, and overdue "
        "remediation items. Any model performing outside acceptable parameters for two consecutive "
        "months is escalated to the BRC for potential retirement or restriction decision."
    )

    save(doc, "Model_Risk_Management_2025")


def doc_21_fraud_prevention():
    doc = Document()
    doc.add_heading("Fraud Prevention and Detection Policy", level=0)
    doc.add_paragraph("Document ID: POL-FRD-2025-021")
    doc.add_paragraph("Effective Date: March 1, 2025")

    doc.add_heading("1. Fraud Risk Assessment", level=1)
    doc.add_paragraph(
        "The Fraud Risk Assessment (FRA) is conducted annually by the Operational Risk Unit in "
        "coordination with the Fraud Prevention Unit (FPU). The FRA maps fraud typologies to business "
        "processes, assigns inherent risk ratings, evaluates control effectiveness, and determines "
        "residual risk. Results are reported to the Operational Risk Committee (ORC) and the BACC."
    )

    doc.add_heading("2. Fraud Detection Systems", level=1)
    doc.add_paragraph(
        "The Bank operates three fraud detection systems: (a) Real-time card fraud detection "
        "(managed by PayTech Solutions per Vendor Risk Management Policy), (b) Online banking fraud "
        "detection (behavioral analytics engine), (c) Internal fraud monitoring (access pattern "
        "anomaly detection per IT Security Policy, Section 5). All systems generate alerts that "
        "are triaged by the FPU within 4 hours."
    )

    doc.add_heading("3. Fraud Investigation", level=1)
    doc.add_paragraph(
        "Confirmed fraud cases are investigated by the FPU in coordination with: Legal Department "
        "(for recovery and prosecution), HR Department (for internal fraud disciplinary action per "
        "HR Policy Section 5), IT Security (for digital forensics per IT Security Policy Section 4.3), "
        "and the MLRO (if fraud indicators overlap with money laundering per AML Policy Section 3.2). "
        "Investigation reports are submitted to the BACC."
    )

    doc.add_heading("4. Fraud Loss Reporting", level=1)
    doc.add_paragraph(
        "Fraud losses are categorized per the Operational Risk Framework loss taxonomy. Losses "
        "exceeding EGP 1 million are individually reported to the ORC. Aggregate fraud losses are "
        "reported quarterly to the BRC. Annual fraud loss data feeds into the capital adequacy "
        "calculation under the Advanced Measurement Approach (AMA) per Capital Adequacy Policy, "
        "Section 3. The Central Bank requires notification of fraud losses exceeding EGP 5 million "
        "within 48 hours."
    )

    doc.add_heading("5. Whistleblower Program", level=1)
    doc.add_paragraph(
        "The Bank maintains an anonymous whistleblower hotline managed by the Compliance Department. "
        "All reports are investigated within 10 business days. Whistleblowers are protected under "
        "the Bank's Non-Retaliation Policy. The CCO reports whistleblower statistics (anonymized) "
        "to the BACC quarterly. Substantiated reports involving senior management are escalated "
        "directly to the Board Chairman."
    )

    save(doc, "Fraud_Prevention_Policy_2025")


def doc_22_outsourcing_policy():
    doc = Document()
    doc.add_heading("Outsourcing and Offshoring Policy", level=0)
    doc.add_paragraph("Document ID: POL-OUT-2025-022")
    doc.add_paragraph("Effective Date: February 1, 2025")

    doc.add_heading("1. Outsourcing Governance", level=1)
    doc.add_paragraph(
        "The Board of Directors retains ultimate accountability for outsourced activities. The "
        "Outsourcing Oversight Committee (OOC), chaired by the COO and including CRO, CTO, and CCO, "
        "approves all material outsourcing arrangements. Material outsourcing includes any activity "
        "that, if disrupted, would materially impact the Bank's operations, reputation, or regulatory "
        "compliance. The OOC reports to the Board Governance Committee quarterly."
    )

    doc.add_heading("2. Regulatory Requirements", level=1)
    doc.add_paragraph(
        "Per Central Bank Circular 2024/OUT/001, the Bank must: (a) notify the Central Bank 60 days "
        "before any material outsourcing, (b) maintain an outsourcing register accessible to supervisors, "
        "(c) ensure outsourced activities meet the same standards as if performed internally, "
        "(d) retain core management and control functions in-house. Cross-border outsourcing requires "
        "additional Board approval and Central Bank notification."
    )

    doc.add_heading("3. Core vs Non-Core Activities", level=1)
    doc.add_paragraph(
        "Core activities (prohibited from outsourcing): strategic decision-making, risk management "
        "oversight, compliance oversight, internal audit, and AML/CTF reporting. Activities with "
        "restricted outsourcing (require Board approval): IT infrastructure management, customer "
        "data processing, core banking system operations, and payment processing. The classification "
        "is reviewed annually by the OOC."
    )

    doc.add_heading("4. Performance Monitoring", level=1)
    doc.add_paragraph(
        "Outsourced activities are monitored using Key Performance Indicators (KPIs) and Key Risk "
        "Indicators (KRIs). SLA breaches are categorized as: Minor (within tolerance), Material "
        "(exceeding tolerance), and Critical (service disruption). Material and Critical SLA breaches "
        "for Tier-1 vendors trigger escalation per the Vendor Risk Management Policy Section 4 "
        "and may activate BCM protocols per the BCM Policy."
    )

    doc.add_heading("5. Exit Strategy", level=1)
    doc.add_paragraph(
        "Every material outsourcing arrangement must have a documented exit strategy. The exit plan "
        "must include: transition timeline (maximum 12 months for critical functions), data migration "
        "procedures (per Data Governance Policy), staff augmentation requirements, and fallback "
        "arrangements. Exit strategy testing is conducted biennially for Tier-1 vendors. Results "
        "are reported to the OOC and the Board Governance Committee."
    )

    save(doc, "Outsourcing_Policy_2025")


def doc_23_ifrs9_policy():
    doc = Document()
    doc.add_heading("IFRS 9 Expected Credit Loss (ECL) Policy", level=0)
    doc.add_paragraph("Document ID: POL-ECL-2025-023")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. ECL Governance", level=1)
    doc.add_paragraph(
        "The ECL governance framework involves: (a) the Board Audit and Compliance Committee (BACC) "
        "approves the ECL methodology, (b) the CFO owns the ECL calculation process, (c) the Model "
        "Risk Committee (per Model Risk Management Policy) validates the ECL models, (d) Internal "
        "Audit provides independent assurance (per Internal Audit Charter). The Credit Risk Unit "
        "supplies the credit data inputs."
    )

    doc.add_heading("2. Staging Criteria", level=1)
    doc.add_paragraph(
        "Stage 1 (12-month ECL): performing assets with no significant increase in credit risk (SICR). "
        "Stage 2 (lifetime ECL): assets with SICR, defined as: 30+ days past due, 2+ notch downgrade "
        "in internal rating, or qualitative indicators (restructured, watch-listed). Stage 3 (credit "
        "impaired): 90+ days past due, or in default per the Credit Risk Policy default definition "
        "(Section 5.1)."
    )

    doc.add_heading("3. ECL Model Components", level=1)
    doc.add_paragraph(
        "ECL = PD × LGD × EAD. Probability of Default (PD) models use logistic regression calibrated "
        "on 5 years of internal data. Loss Given Default (LGD) uses recovery rate analysis. Exposure "
        "at Default (EAD) uses credit conversion factors per the Capital Adequacy Policy. Forward-looking "
        "adjustments incorporate three macroeconomic scenarios: baseline (50% weight), upside (25%), "
        "and downside (25%). GDP growth, unemployment rate, and interest rates are the key macro variables."
    )

    doc.add_heading("4. Significant Judgments", level=1)
    doc.add_paragraph(
        "Key judgments requiring senior management approval: (a) SICR threshold calibration — reviewed "
        "semi-annually by the Credit Risk Committee, (b) Forward-looking scenario weights — approved "
        "by ALCO, (c) Post-model adjustments (management overlays) — individually justified and approved "
        "by CFO for amounts below EGP 50 million and by the BACC for larger amounts, (d) Sector-wide "
        "staging overrides — approved by the Credit Risk Committee and reported to the BRC."
    )

    doc.add_heading("5. ECL Reporting", level=1)
    doc.add_paragraph(
        "Monthly ECL calculations are reviewed by the CFO and Head of Risk Analytics. Quarterly ECL "
        "results are presented to the BACC with: stage migration analysis, top 20 individual ECL "
        "exposures, sector concentration analysis, and macro scenario sensitivity. External auditors "
        "review the ECL calculation semi-annually. IFRS 9 disclosures in financial statements follow "
        "the Central Bank disclosure template."
    )

    save(doc, "IFRS9_ECL_Policy_2025")


def doc_24_related_party():
    doc = Document()
    doc.add_heading("Related Party Transactions Policy", level=0)
    doc.add_paragraph("Document ID: POL-RPT-2025-024")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Definition of Related Parties", level=1)
    doc.add_paragraph(
        "Related parties include: Board members and their immediate family (spouse, parents, children), "
        "senior management (C-suite and department heads), major shareholders (holding 5% or more of "
        "voting shares), subsidiaries and affiliates, and entities where the Bank's Board members or "
        "senior management hold significant influence. The Compliance Department maintains the Related "
        "Party Register, updated quarterly."
    )

    doc.add_heading("2. Transaction Limits", level=1)
    doc.add_paragraph(
        "Aggregate exposure to all related parties must not exceed 20% of the Bank's Tier-1 capital. "
        "Individual related party exposure must not exceed 5% of Tier-1 capital. These limits align "
        "with Central Bank large exposure regulations and the Capital Adequacy Policy, Section 5.2. "
        "Credit facilities to related parties must be on arms-length terms as verified by the "
        "Credit Risk Unit."
    )

    doc.add_heading("3. Approval Requirements", level=1)
    doc.add_paragraph(
        "All related party transactions require: (a) independent director approval (majority of "
        "independent Board members), (b) competitive market terms verification by the relevant "
        "business unit, (c) Compliance Department confirmation that the transaction is within limits, "
        "(d) Internal Audit post-transaction review within 30 days. Board members must recuse "
        "themselves from voting on transactions involving them personally."
    )

    doc.add_heading("4. Disclosure Requirements", level=1)
    doc.add_paragraph(
        "Related party transactions are disclosed: (a) in the Bank's annual financial statements "
        "per IAS 24, (b) in regulatory returns to the Central Bank (quarterly), (c) to the BACC "
        "at each meeting. The CCO provides an annual Related Party Transaction Summary to the "
        "Board Governance Committee, including any new relationships, limit utilization, and "
        "compliance exceptions."
    )

    doc.add_heading("5. Monitoring and Reporting", level=1)
    doc.add_paragraph(
        "The Credit Risk Unit flags any credit facility that approaches 80% of the individual "
        "related party limit. The Treasury Department flags any investment or trading transaction "
        "with related party counterparties to the CCO before execution. The Internal Audit "
        "Charter (Section 3) includes annual review of related party transaction compliance "
        "as a mandatory audit topic."
    )

    save(doc, "Related_Party_Policy_2025")


def doc_25_climate_risk():
    doc = Document()
    doc.add_heading("Climate and Environmental Risk Management Framework", level=0)
    doc.add_paragraph("Document ID: FRM-CLM-2025-025")
    doc.add_paragraph("Effective Date: March 1, 2025")

    doc.add_heading("1. Governance of Climate Risk", level=1)
    doc.add_paragraph(
        "Climate risk oversight is assigned to the Board Risk Committee (BRC). The CRO designates "
        "a Climate Risk Officer within the Risk Management Department. Climate risk is integrated "
        "into the existing Risk Management Framework (credit, market, operational) rather than "
        "treated as a separate risk category. The Climate Risk Officer reports quarterly to the BRC "
        "and annually to the full Board."
    )

    doc.add_heading("2. Physical Risk Assessment", level=1)
    doc.add_paragraph(
        "Physical risks (flooding, heat stress, water scarcity) are assessed for the Bank's: "
        "(a) own premises — BCM Policy Section 1 recovery objectives apply, (b) collateral portfolio "
        "— Credit Risk Policy Section 6 collateral valuation must incorporate physical risk discounts "
        "for properties in high-risk zones, (c) borrower operations — sector risk assessment per "
        "Credit Risk Policy Section 7 includes climate vulnerability scoring."
    )

    doc.add_heading("3. Transition Risk Assessment", level=1)
    doc.add_paragraph(
        "Transition risks from policy changes, technology shifts, and market repricing are assessed "
        "through: (a) Carbon-intensive sector identification — real estate, transportation, "
        "manufacturing, oil and gas, (b) Portfolio carbon footprint estimation, (c) Stress testing "
        "using transition scenarios from the Network for Greening the Financial System (NGFS). "
        "Results feed into the IFRS 9 ECL forward-looking overlay per the IFRS 9 Policy, Section 4."
    )

    doc.add_heading("4. Green Lending Framework", level=1)
    doc.add_paragraph(
        "The Bank offers preferential terms for: (a) Energy efficiency projects — 50bps rate discount, "
        "(b) Renewable energy financing — extended tenor up to 15 years, (c) Green building "
        "certification — reduced collateral requirements. Green lending targets are set by ALCO "
        "and monitored by the Board Risk Committee. The Credit Risk Committee reviews sector "
        "limits annually with climate risk adjustments."
    )

    doc.add_heading("5. TCFD Disclosure", level=1)
    doc.add_paragraph(
        "The Bank publishes annual climate risk disclosures following the Task Force on Climate-related "
        "Financial Disclosures (TCFD) framework. Disclosures cover: governance arrangements, strategy "
        "and scenario analysis, risk management integration, and metrics/targets. The disclosure "
        "is reviewed by the BACC and approved by the Board of Directors before publication."
    )

    save(doc, "Climate_Risk_Framework_2025")


def doc_26_digital_banking():
    doc = Document()
    doc.add_heading("Digital Banking and API Policy", level=0)
    doc.add_paragraph("Document ID: POL-DIG-2025-026")
    doc.add_paragraph("Effective Date: February 1, 2025")

    doc.add_heading("1. Digital Channel Governance", level=1)
    doc.add_paragraph(
        "Digital banking channels (mobile app, internet banking, Open Banking APIs) are governed by "
        "the Digital Banking Committee (DBC), chaired by the Chief Digital Officer (CDO) with members "
        "from IT, Risk, Compliance, and Customer Experience. The DBC reports to the Board IT Committee. "
        "All digital product launches require DBC approval after: security assessment (per IT Security "
        "Policy), data privacy review (per Data Governance Policy), and consumer protection check "
        "(per Consumer Protection Policy)."
    )

    doc.add_heading("2. Open Banking and API Management", level=1)
    doc.add_paragraph(
        "The Bank's Open Banking APIs comply with Central Bank Open Banking Regulation 2024/OB/001. "
        "Third-party providers (TPPs) must be licensed by the Central Bank. API access is governed by "
        "OAuth 2.0 with mutual TLS authentication. Rate limits, consent management, and data minimization "
        "principles are enforced at the API gateway. TPP onboarding follows the Vendor Risk Management "
        "Policy with automatic Tier-2 classification."
    )

    doc.add_heading("3. Digital Identity and Authentication", level=1)
    doc.add_paragraph(
        "Customer authentication for digital channels uses: (a) Level 1 (informational access) — "
        "username + password + device fingerprint, (b) Level 2 (transactions below EGP 50,000) — "
        "Level 1 + OTP, (c) Level 3 (transactions above EGP 50,000 or sensitive operations) — "
        "Level 2 + biometric verification. Biometric modalities include fingerprint, facial recognition, "
        "and voice recognition per the IT Security Policy Section 3.2."
    )

    doc.add_heading("4. Digital Lending", level=1)
    doc.add_paragraph(
        "Pre-approved digital loans up to EGP 500,000 may be originated through mobile and internet "
        "banking. The credit decision uses the automated credit scoring model (per Model Risk "
        "Management Policy). Loans are subject to the same credit risk policies as branch-originated "
        "loans (per Credit Risk Policy). Digital loan disbursement requires real-time fraud check "
        "(per Fraud Prevention Policy, Section 2)."
    )

    doc.add_heading("5. Digital Customer Onboarding", level=1)
    doc.add_paragraph(
        "Video-KYC for remote account opening is permitted per Central Bank Circular 2024/KYC/002. "
        "The process includes: live video interview with trained staff, OCR of ID documents, liveness "
        "detection, and screening against AML watchlists (per AML Policy, Section 4). Digital CDD "
        "records are retained per the Data Governance Policy retention schedule and AML Policy "
        "Section 5 record retention requirements."
    )

    save(doc, "Digital_Banking_Policy_2025")


def doc_27_stress_testing():
    doc = Document()
    doc.add_heading("Enterprise Stress Testing Framework", level=0)
    doc.add_paragraph("Document ID: FRM-STR-2025-027")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Stress Testing Governance", level=1)
    doc.add_paragraph(
        "The Enterprise Stress Testing Program is overseen by the CRO and governed by the Board Risk "
        "Committee (BRC). The Risk Analytics Unit designs scenarios and executes stress tests. Results "
        "inform: (a) capital planning per Capital Adequacy Policy, (b) liquidity contingency planning "
        "per Liquidity Policy, (c) risk appetite calibration, (d) recovery plan triggers. "
        "The Central Bank mandates annual stress testing submissions per Circular 2024/ST/001."
    )

    doc.add_heading("2. Scenario Design", level=1)
    doc.add_paragraph(
        "Three scenario types are maintained: (a) Historical — replication of 2008 financial crisis, "
        "2011 Egyptian revolution, 2016 devaluation, 2023 currency crisis, (b) Hypothetical — designed "
        "around emerging risks (pandemic, cyberattack, geopolitical escalation), (c) Reverse stress "
        "testing — identifies scenarios that would render the Bank non-viable (capital ratio below "
        "the regulatory minimum of 12.5% per Capital Adequacy Policy, Section 2)."
    )

    doc.add_heading("3. Impact Assessment", level=1)
    doc.add_paragraph(
        "Stress test impacts are measured across: (a) Credit losses — PD and LGD shocks applied to "
        "the loan portfolio using IFRS 9 ECL models (per IFRS 9 Policy), (b) Market losses — VaR "
        "and stressed VaR recalculation (per Market Risk Framework), (c) Operational losses — "
        "scenario-based estimation using Operational Risk Framework methodology, (d) Liquidity "
        "impact — projected outflows and HQLA haircuts per Liquidity Policy stress parameters."
    )

    doc.add_heading("4. Recovery Plan Integration", level=1)
    doc.add_paragraph(
        "The Bank's Recovery Plan identifies recovery actions to be taken when stress test indicators "
        "breach early warning triggers: (a) Capital ratio below 14% (regulatory minimum is 12.5%) — "
        "ALCO convenes within 48 hours to assess capital actions, (b) LCR below 110% — Treasury "
        "activates the Contingency Funding Plan (per Liquidity Policy Section 5), (c) CET1 ratio "
        "below 10% — Board of Directors emergency meeting within 24 hours to approve capital "
        "restoration plan. Recovery Plan is tested annually via tabletop exercise."
    )

    doc.add_heading("5. Reporting and Communication", level=1)
    doc.add_paragraph(
        "Stress test results are: (a) presented to the BRC within 30 days of completion, (b) submitted "
        "to the Central Bank per regulatory deadline, (c) shared with ALCO for capital and liquidity "
        "planning, (d) included in the Annual Risk Report to the Board. Key stress test metrics are: "
        "peak-to-trough capital ratio, maximum liquidity surplus/deficit, projected net income impact, "
        "and time to recovery."
    )

    save(doc, "Stress_Testing_Framework_2025")


def doc_28_sanctions_policy():
    doc = Document()
    doc.add_heading("Sanctions Compliance Policy", level=0)
    doc.add_paragraph("Document ID: POL-SAN-2025-028")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Sanctions Program Governance", level=1)
    doc.add_paragraph(
        "The Sanctions Compliance Program is managed by the Sanctions Officer who reports to the CCO. "
        "The Compliance Committee provides strategic direction. The Board Audit and Compliance Committee "
        "(BACC) approves the annual Sanctions Risk Assessment. The Bank maintains zero tolerance for "
        "sanctions violations. Any intentional circumvention is grounds for immediate termination per "
        "the HR Policy disciplinary framework."
    )

    doc.add_heading("2. Screening Requirements", level=1)
    doc.add_paragraph(
        "Screening is performed against: UN Security Council Consolidated List, OFAC SDN and Sectoral "
        "Sanctions lists, EU Consolidated Sanctions List, Central Bank of Egypt restricted list, and "
        "any additional lists designated by the Compliance Committee. Screening triggers: customer "
        "onboarding, all wire transfers (SWIFT MT103, MT202), periodic batch re-screening (monthly "
        "for all customers), and any name/address change in customer records."
    )

    doc.add_heading("3. Alert Handling", level=1)
    doc.add_paragraph(
        "Screening alerts are categorized as: Exact Match (automatic freeze), Close Match (manual review "
        "within 4 hours), and Possible Match (manual review within 24 hours). Alert disposition must "
        "be documented with rationale. The Sanctions Officer reviews all Exact and Close Match "
        "dispositions. Monthly alert statistics are reported to the Compliance Committee. "
        "False-positive rates exceeding 98% trigger system calibration review by the vendor "
        "(per Vendor Risk Management Policy, Tier-1 monitoring requirements)."
    )

    doc.add_heading("4. Sanctions Breach Procedures", level=1)
    doc.add_paragraph(
        "In the event of a confirmed sanctions breach: (a) the transaction is immediately frozen, "
        "(b) the Sanctions Officer notifies the CCO within 1 hour, (c) the CCO notifies the CEO "
        "and Legal Department within 2 hours, (d) regulatory notification to the Central Bank "
        "and relevant sanctions authority within 24 hours, (e) the Board Chairman is informed "
        "within 48 hours. External legal counsel is engaged immediately for OFAC-related breaches. "
        "All sanctions breaches are reportable operational risk events per the Operational Risk "
        "Framework, Section 3."
    )

    doc.add_heading("5. Correspondent Banking Sanctions Risk", level=1)
    doc.add_paragraph(
        "Correspondent banking relationships receive enhanced sanctions scrutiny. Downstream correspondent "
        "(nested) risk is assessed during onboarding and annually (per AML Policy Section 2.2 and "
        "Payment Systems Policy Section 5). The Bank does not process payable-through account "
        "transactions. SWIFT MT202COV messages require both originator and beneficiary screening."
    )

    save(doc, "Sanctions_Compliance_Policy_2025")


def doc_29_remuneration():
    doc = Document()
    doc.add_heading("Remuneration and Incentive Policy", level=0)
    doc.add_paragraph("Document ID: POL-REM-2025-029")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Remuneration Governance", level=1)
    doc.add_paragraph(
        "The Board Remuneration Committee (BRemC) oversees the Bank's remuneration framework. BRemC "
        "membership comprises at least 3 independent non-executive directors. The committee meets "
        "quarterly and reports to the Board of Directors. The CRO and CCO attend BRemC meetings as "
        "permanent invitees to provide risk and compliance input on incentive structures."
    )

    doc.add_heading("2. Risk-Adjusted Compensation", level=1)
    doc.add_paragraph(
        "Variable compensation for Material Risk Takers (MRTs) — including all C-suite, department "
        "heads, and trading desk senior dealers — is subject to: (a) maximum variable-to-fixed ratio "
        "of 200%, (b) at least 40% deferral over 3 years (60% for senior executives), (c) at least "
        "50% of deferred variable paid in equity-linked instruments. MRTs are identified annually "
        "by the CRO based on quantitative and qualitative criteria per Central Bank guidance."
    )

    doc.add_heading("3. Malus and Clawback", level=1)
    doc.add_paragraph(
        "Deferred compensation is subject to malus (reduction) and clawback (recovery) provisions. "
        "Trigger events include: (a) material financial restatement affecting the unit's performance, "
        "(b) individual misconduct (fraud, compliance violations per HR Policy Section 5), "
        "(c) risk management failures resulting in losses exceeding the unit's risk appetite, "
        "(d) regulatory enforcement action attributable to the individual. The BRemC decides on "
        "malus/clawback activation with CRO and CCO input."
    )

    doc.add_heading("4. Trading Desk Incentives", level=1)
    doc.add_paragraph(
        "FX Trading Desk, Fixed Income, and Equities desk compensation is based on: (a) 60% risk-adjusted "
        "P&L (net of VaR-based capital charges per Market Risk Framework), (b) 20% compliance and "
        "control metrics (limit breaches, trade errors, regulatory findings), (c) 20% qualitative "
        "assessment (teamwork, knowledge sharing, client relationship). Desk heads certify that "
        "P&L figures used for bonus calculations are reconciled with Finance Department records."
    )

    doc.add_heading("5. Disclosure", level=1)
    doc.add_paragraph(
        "The Bank discloses remuneration information in the annual report per Central Bank Pillar 3 "
        "requirements and following the Capital Adequacy Policy disclosure schedule. Disclosures "
        "include: aggregate remuneration by business area, number of MRTs, deferred remuneration "
        "details, and any malus/clawback activations during the year. The BRemC Chair presents "
        "the Remuneration Report to shareholders at the Annual General Meeting."
    )

    save(doc, "Remuneration_Policy_2025")


def doc_30_recovery_resolution():
    doc = Document()
    doc.add_heading("Recovery and Resolution Plan", level=0)
    doc.add_paragraph("Document ID: PLN-RRP-2025-030")
    doc.add_paragraph("Effective Date: January 1, 2025")

    doc.add_heading("1. Recovery Plan Overview", level=1)
    doc.add_paragraph(
        "The Recovery Plan is a Board-approved document identifying credible actions the Bank can take "
        "to restore financial health during severe stress without public support. The plan is updated "
        "annually under the coordination of the CRO and CFO, with input from Treasury (liquidity "
        "recovery actions), Credit Risk (asset disposal options), and Legal (regulatory obligations). "
        "The Board Risk Committee reviews and the Board of Directors approves the plan."
    )

    doc.add_heading("2. Recovery Indicators and Triggers", level=1)
    doc.add_paragraph(
        "Recovery plan activation is based on quantitative indicators: (a) CET1 ratio below 10.5% "
        "(normal operating minimum per Capital Adequacy Policy: 12.5%), (b) Total Capital Ratio "
        "below 14%, (c) LCR below 100% for 3 consecutive days (per Liquidity Policy), (d) NSFR "
        "below 100%, (e) Single-day deposit outflow exceeding 5% of total deposits, (f) Credit "
        "rating downgrade by 2+ notches. Each trigger has an owner: CRO monitors capital triggers, "
        "CFO/Treasury monitors liquidity triggers."
    )

    doc.add_heading("3. Recovery Actions", level=1)
    doc.add_heading("3.1 Capital Recovery", level=2)
    doc.add_paragraph(
        "Capital recovery actions in order of severity: (a) Suspend dividend payments (Board decision, "
        "notify Central Bank), (b) Activate AT1 instrument terms (conversion or write-down per Capital "
        "Adequacy Policy Section 4), (c) Sell non-core subsidiaries or investment portfolio assets "
        "(ALCO authorization, execution by Treasury), (d) Emergency rights issue (Board and "
        "shareholder approval, 90-day timeline), (e) Asset portfolio run-off — cease new lending "
        "in non-priority segments (Board Credit Committee decision)."
    )

    doc.add_heading("3.2 Liquidity Recovery", level=2)
    doc.add_paragraph(
        "Liquidity recovery actions: (a) Activate Contingency Funding Plan (per Liquidity Policy "
        "Section 5), (b) Access Central Bank lending facilities — Emergency Liquidity Assistance "
        "(ELA) with eligible collateral, (c) Conduct emergency repo operations on HQLA portfolio "
        "(Treasury execution), (d) Restrict large withdrawals above EGP 5 million to 24-hour "
        "notice (regulatory approval required), (e) Interbank line drawdown from pre-arranged "
        "committed facilities."
    )

    doc.add_heading("4. Communication Plan", level=1)
    doc.add_paragraph(
        "Recovery plan activation communication: (a) Board Chairman informed immediately, (b) Central "
        "Bank Governor's office notified within 6 hours by the CEO, (c) External auditors notified "
        "within 24 hours by the CFO, (d) Public communication (if required) managed by Corporate "
        "Communications with Board and Central Bank pre-approval, (e) Staff communication via "
        "internal memo by HR per the HR Policy crisis communication module."
    )

    doc.add_heading("5. Resolution Planning", level=1)
    doc.add_paragraph(
        "The Resolution Plan is maintained in coordination with the Central Bank's Resolution Authority. "
        "Key components: (a) Critical functions mapping (aligned with BCM Policy business impact "
        "analysis), (b) Bail-in eligible liabilities identification per Capital Adequacy Policy, "
        "(c) Separability analysis for subsidiaries, (d) Access to Financial Market Infrastructures "
        "(RTGS, SWIFT, ACH per Payment Systems Policy), (e) MIS capabilities to provide resolution "
        "authority with required data within 24 hours."
    )

    save(doc, "Recovery_Resolution_Plan_2025")


if __name__ == "__main__":
    print("Creating 20 additional banking documents (docs 11-30)...")
    print()

    doc_11_treasury_operations()
    doc_12_compliance_monitoring()
    doc_13_hr_policy()
    doc_14_vendor_management()
    doc_15_bcm_policy()
    doc_16_investment_policy()
    doc_17_consumer_protection()
    doc_18_market_risk()
    doc_19_payment_systems()
    doc_20_model_risk()
    doc_21_fraud_prevention()
    doc_22_outsourcing_policy()
    doc_23_ifrs9_policy()
    doc_24_related_party()
    doc_25_climate_risk()
    doc_26_digital_banking()
    doc_27_stress_testing()
    doc_28_sanctions_policy()
    doc_29_remuneration()
    doc_30_recovery_resolution()

    print()
    print(f"Done! Created 20 documents in {OUTPUT_DIR}")
    print("Run: python build_index.py  to rebuild the index.")
